"""
Convert AWS_SETUP_INSTRUCTIONS.md and AWS_COMMANDS_AND_CONSOLE.md to .docx
with professional formatting. Preserves content word-for-word.
Run from repo root: python scripts/md_to_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill):
    """Set table cell background color (hex without #)."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph(doc, text, style=None, paragraph=None):
    """Add paragraph (or add runs to existing paragraph), handling **bold** and `code`."""
    if not text.strip():
        return doc.add_paragraph() if paragraph is None else paragraph
    if paragraph is None:
        paragraph = doc.add_paragraph(style=style)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
    rest = text
    while rest:
        m = re.search(r"\*\*([^*]+)\*\*", rest)
        mc = re.search(r"`([^`]+)`", rest)
        pos_b = m.start() if m else len(rest)
        pos_c = mc.start() if mc else len(rest)
        if pos_b == len(rest) and pos_c == len(rest):
            paragraph.add_run(rest)
            break
        if pos_b <= pos_c and m:
            paragraph.add_run(rest[: m.start()])
            r = paragraph.add_run(m.group(1))
            r.bold = True
            rest = rest[m.end() :]
        elif mc:
            paragraph.add_run(rest[: mc.start()])
            r = paragraph.add_run(mc.group(1))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            rest = rest[mc.end() :]
        else:
            paragraph.add_run(rest)
            break
    return paragraph


def add_code_block(doc, code, style_name="Normal"):
    """Add a code block with monospace and tight spacing."""
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def parse_md_table(lines):
    """Parse markdown table into list of rows (list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if re.match(r"^[-:\s|]+$", line):
            continue  # skip separator row
        rows.append(cells)
    return rows


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    """Convert one markdown file to docx with professional formatting."""
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Slightly wider than default margins
    section = doc.sections[0]
    margin = Inches(0.5)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin

    # Set default font and tight spacing for all styles
    def set_tight_paragraph_format(style, space_before=0, space_after=6):
        pf = style.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = Pt(13)  # slightly more than 11pt font

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    set_tight_paragraph_format(style, space_before=0, space_after=4)

    # Heading 1: title style, minimal space above/below
    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = "Calibri Light"
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    set_tight_paragraph_format(h1_style, space_before=12, space_after=6)
    # Heading 2
    h2_style = doc.styles["Heading 2"]
    h2_style.font.name = "Calibri"
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    set_tight_paragraph_format(h2_style, space_before=10, space_after=4)
    # Heading 3
    h3_style = doc.styles["Heading 3"]
    h3_style.font.name = "Calibri"
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    set_tight_paragraph_format(h3_style, space_before=8, space_after=4)
    # List styles
    for list_style in ("List Bullet", "List Number"):
        if list_style in doc.styles:
            set_tight_paragraph_format(doc.styles[list_style], space_before=0, space_after=2)

    # Collapse consecutive blank lines to one (reduces excess spacing in output)
    raw_lines = text.split("\n")
    lines = []
    prev_blank = False
    for L in raw_lines:
        is_blank = not L.strip()
        if is_blank and prev_blank:
            continue
        lines.append(L)
        prev_blank = is_blank
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    in_toc_section = False  # output TOC as bullet list (direct copy of .md TOC)

    while i < len(lines):
        line = lines[i]
        raw = line

        # After "## Table of contents": output each TOC line as a bullet with link text only
        if in_toc_section:
            if not line.strip():
                i += 1
                continue
            toc_match = re.match(r"^\s*-\s+\[(.+)\]\(#.+\)", line)
            if toc_match:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                add_paragraph(doc, toc_match.group(1).strip(), paragraph=p)
                i += 1
                continue
            in_toc_section = False
            # Fall through to process this line (e.g. --- or ##)

        # Code block
        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                add_code_block(doc, code_text)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table: collect consecutive table rows
        if line.strip().startswith("|") and "|" in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                rows = parse_md_table(table_lines)
                if rows:
                    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    t.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, cell_text in enumerate(row):
                            if ci < len(t.rows[ri].cells):
                                cell = t.rows[ri].cells[ci]
                                cell.text = cell_text
                                if ri == 0:
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.bold = True
                    # Small spacer after table instead of full blank paragraph
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(2)
                    spacer.paragraph_format.space_after = Pt(4)
                in_table = False
                table_lines = []
            in_table = False

        # Horizontal rule: minimal section break
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            if heading_text == "Table of contents":
                in_toc_section = True
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # List items
        if re.match(r"^[-*] ", line) or re.match(r"^\d+\. ", line):
            rest = line.strip()
            if rest[0] in "-*":
                rest = rest[1:].strip()
            elif re.match(r"^\d+\.", rest):
                rest = re.sub(r"^\d+\.\s*", "", rest)
            p = doc.add_paragraph(style="List Bullet" if re.match(r"^[-*] ", line) else "List Number")
            add_paragraph(doc, rest, paragraph=p)
            i += 1
            continue

        # Empty line: skip (no extra spacer; spacing from --- and paragraph styles only)
        if not line.strip():
            i += 1
            continue

        # Normal paragraph (may contain **bold** and `code`)
        add_paragraph(doc, line)
        i += 1
        continue

    # Flush remaining table
    if in_table and table_lines:
        rows = parse_md_table(table_lines)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci < len(t.rows[ri].cells):
                        t.rows[ri].cells[ci].text = cell_text
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)

    doc.save(docx_path)
    print(f"Saved {docx_path}")


def main():
    aws_setup = Path(__file__).parent.parent / "aws_setup"
    md_to_docx(aws_setup / "AWS_SETUP_INSTRUCTIONS.md", aws_setup / "AWS_SETUP_INSTRUCTIONS.docx")
    md_to_docx(aws_setup / "AWS_COMMANDS_AND_CONSOLE.md", aws_setup / "AWS_COMMANDS_AND_CONSOLE.docx")


if __name__ == "__main__":
    main()
