"""Audit the current color-coded final report DOCX.

This audit matches the June 28 final-report workflow: the uploaded source DOCX
is patched in our sections, then the color-coded DOCX marks teammate content
gray and our GivingTuesday/Q9 content black.
"""

from __future__ import annotations

import os
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


REPO_ROOT = Path(__file__).resolve().parents[2]
DOCX = Path(
    os.environ.get(
        "FINAL_REPORT_DOCX",
        REPO_ROOT
        / "docs"
        / "deliverables"
        / "final_report"
        / "working"
        / "black_hills_philanthropy_color_coded.docx",
    )
)

GRAY = "666666"
BLACK = "000000"
ADDED_TABLE_HEADER_FILL = "DEEAF0"
REGION_ORDER = ["Black Hills", "Billings", "Flagstaff", "Missoula", "Sioux Falls"]
NEW_MEDIAN_GAP_HEADER = "Black Hills median minus benchmark median (95% bootstrap CI)"
TABLE_FONT_NAME = "Arial"
TABLE_FONT_SIZE_PT = 8.0
FIGURE_TITLE_FONT_SIZE_PT = 10.0

MAIN_FIGURES = [
    "Figure 5-2: Total Revenue by Region",
    "Figure 5-3: Program Service Revenue by Region",
    "Figure 5-4: Total Contributions by Region",
    "Figure 5-5: Government Grants Received by Region",
    "Figure 5-9: Fundraising Event Contributions by Region",
]
APPENDIX_ONLY_FIGURES = [
    "Figure 5-6: Federated Campaign Contributions by Region",
    "Figure 5-7: Related Organization Contributions by Region",
    "Figure 5-8: Membership Dues by Region",
    "Figure 5-10: Mixed or Unclassified Contributions by Region",
    "Figure 5-11: Other Revenue by Region",
]
APPENDIX_FIGURES = [
    "Figure A-1: Total Revenue by Region",
    "Figure A-2: Program Service Revenue by Region",
    "Figure A-3: Total Contributions by Region",
    "Figure A-4: Government Grants Received by Region",
    "Figure A-5: Federated Campaign Contributions by Region",
    "Figure A-6: Related Organization Contributions by Region",
    "Figure A-7: Membership Dues by Region",
    "Figure A-8: Fundraising Event Contributions by Region",
    "Figure A-9: Mixed or Unclassified Contributions by Region",
    "Figure A-10: Other Revenue by Region",
]
SOURCE_TABLE_CAPTIONS = [f"Table A-{index}:" for index in range(2, 12)]
Q9_TABLE_TITLES = [
    "Table 5-2a: Median Reported Dollars for Total Revenue by Region",
    "Table 5-3a: Median Reported Dollars for Program Service Revenue by Region",
    "Table 5-4a: Median Reported Dollars for Total Contributions by Region",
    "Table 5-5a: Median Reported Dollars for Government Grants Received by Region",
    "Table 5-9a: Median Reported Dollars for Fundraising Event Contributions by Region",
    "Table A-1: Total Reported Dollars by Revenue Source and Region (Aggregate)",
    "Table A-2: Total Revenue by Region",
    "Table A-3: Program Service Revenue by Region",
    "Table A-4: Total Contributions by Region",
    "Table A-5: Government Grants Received by Region",
    "Table A-6: Federated Campaign Contributions by Region",
    "Table A-7: Related Organization Contributions by Region",
    "Table A-8: Membership Dues by Region",
    "Table A-9: Fundraising Event Contributions by Region",
    "Table A-10: Mixed or Unclassified Contributions by Region",
    "Table A-11: Other Revenue by Region",
    "Table A-12: Revenue Source Definitions",
]


def body_items(document: Document):
    for element in document.element.body.iterchildren():
        if element.tag.endswith("}p"):
            yield Paragraph(element, document._body)
        elif element.tag.endswith("}tbl"):
            yield Table(element, document._body)


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("a:blip")))


def run_color(run) -> str | None:
    if run.font.color.rgb is not None:
        return str(run.font.color.rgb).upper()
    run_properties = run._element.find(qn("w:rPr"))
    color_element = run_properties.find(qn("w:color")) if run_properties is not None else None
    if color_element is not None:
        return (color_element.get(qn("w:val")) or "").upper()
    return None


def paragraph_colors(paragraph: Paragraph) -> list[str]:
    return [color for run in paragraph.runs if run.text for color in [run_color(run)] if color]


def paragraph_is_color(paragraph: Paragraph, color: str) -> bool:
    colors = paragraph_colors(paragraph)
    return bool(colors) and all(item == color for item in colors)


def table_header(table: Table) -> list[str]:
    return [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []


def is_q9_added_table(table: Table) -> bool:
    return table_header(table) in [
        ["Region", "Median reported dollars"],
        [
            "Region",
            "Reporting rate",
            "Positive reporters (n)",
            "Median reported dollars",
            NEW_MEDIAN_GAP_HEADER,
            "Pairwise p-value vs Black Hills",
        ],
        ["Revenue source", "Black Hills", "Billings", "Flagstaff", "Missoula", "Sioux Falls"],
        ["Revenue source", "What it measures", "How it is reported on IRS filings"],
    ]


def section_text(paragraphs: list[Paragraph], start_text: str, end_text: str) -> str:
    start = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == start_text)
    end = next(i for i, paragraph in enumerate(paragraphs[start + 1 :], start=start + 1) if paragraph.text.strip() == end_text)
    return "\n".join(paragraph.text for paragraph in paragraphs[start:end])


def paragraph_index(paragraphs: list[Paragraph], text: str) -> int:
    return next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == text)


def table_after_caption(document: Document, caption: str) -> Table | None:
    seen = False
    for item in body_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text == caption:
                seen = True
                continue
            if seen and text:
                return None
        elif seen:
            return item
    return None


def image_follows_caption(document: Document, caption: str) -> bool:
    seen = False
    for item in body_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if text == caption:
                seen = True
                continue
            if seen:
                if paragraph_has_image(item):
                    return True
                if text:
                    return False
        elif seen:
            return False
    return False


def check_color(document: Document, text: str, color: str, issues: list[str]) -> None:
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == text]
    if len(matches) != 1:
        issues.append(f"Paragraph count for {text!r}: {len(matches)}")
        return
    if not paragraph_is_color(matches[0], color):
        issues.append(f"Unexpected color for {text!r}: {paragraph_colors(matches[0])}")


def main() -> None:
    issues: list[str] = []
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    with ZipFile(DOCX) as zip_file:
        bad = zip_file.testzip()
        if bad is not None:
            issues.append(f"Corrupt DOCX package member: {bad}")

    document = Document(DOCX)
    paragraphs = document.paragraphs
    full = "\n".join(paragraph.text for paragraph in paragraphs)

    if len(document.tables) != 44:
        issues.append(f"Table count: {len(document.tables)} (expect 44)")
    if len(document.inline_shapes) != 51:
        issues.append(f"Inline image count: {len(document.inline_shapes)} (expect 51)")

    expected_heading_once = [
        "IRS Exempt Organizations Business Master File (IRS EO BMF)",
        "GivingTuesday 990 DataMart",
        "National Center for Charitable Statistics (NCCS) Business Master File (BMF)",
        "Appendix A: Non-profit revenue source comparison details",
    ]
    for text in expected_heading_once:
        count = sum(1 for paragraph in paragraphs if paragraph.text.strip() == text)
        if count != 1:
            issues.append(f"Heading count for {text!r}: {count} (expect 1)")

    expected_text_once = [
        "Table A-1: Total Reported Dollars by Revenue Source and Region (Aggregate)",
        "Table A-12: Revenue Source Definitions",
        *APPENDIX_FIGURES,
    ]
    for text in expected_text_once:
        count = full.count(text)
        if count != 1:
            issues.append(f"Text count for {text!r}: {count} (expect 1)")
    for text in MAIN_FIGURES:
        count = full.count(text)
        if count != 1:
            issues.append(f"Text count for main figure {text!r}: {count} (expect 1)")

    for title in [*MAIN_FIGURES, *APPENDIX_FIGURES]:
        matches = [paragraph for paragraph in paragraphs if paragraph.text.strip() == title]
        for paragraph in matches:
            if paragraph.style.name != "Table/Figure Title":
                issues.append(f"Figure title does not use Table/Figure Title style: {title}")
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                issues.append(f"Figure title is not left-aligned: {title}")
            if not all(run.bold for run in paragraph.runs if run.text):
                issues.append(f"Figure title is not bold: {title}")
            if any(run.underline for run in paragraph.runs if run.text):
                issues.append(f"Figure title is underlined: {title}")
            for run in paragraph.runs:
                if not run.text:
                    continue
                if run.font.name != TABLE_FONT_NAME:
                    issues.append(f"Figure title font is not Arial: {title}")
                size = run.font.size.pt if run.font.size is not None else None
                if size != FIGURE_TITLE_FONT_SIZE_PT:
                    issues.append(f"Figure title font size is not 10 pt: {title} ({size!r})")
        if not image_follows_caption(document, title):
            issues.append(f"Figure title is not directly above an image: {title}")

    for title in Q9_TABLE_TITLES:
        matches = [paragraph for paragraph in paragraphs if paragraph.text.strip() == title]
        expected_count = 1
        if len(matches) != expected_count:
            issues.append(f"Table title count for {title!r}: {len(matches)} (expect {expected_count})")
        for paragraph in matches:
            if paragraph.style.name != "Table/Figure Title":
                issues.append(f"Table title does not use Table/Figure Title style: {title}")
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                issues.append(f"Table title is not left-aligned: {title}")
            if not all(run.bold for run in paragraph.runs if run.text):
                issues.append(f"Table title is not bold: {title}")
            if any(run.underline for run in paragraph.runs if run.text):
                issues.append(f"Table title is underlined: {title}")
        if table_after_caption(document, title) is None:
            issues.append(f"Table title is not directly above a table: {title}")

    for stale in [
        "Add content here",
        "(coverage limits, form types, and what is not in the file)",
        "Not applicable (baseline)",
        "Reference",
        "Figure 5-1:",
        "Black Hills minus benchmark region (95% CI)",
        "Table 5-1: Median Reported Dollars by Revenue Source and Region",
        "Table 5-1 provides the exact median reported dollars",
        "summarized in Table 5-1",
    ]:
        if stale in full:
            issues.append(f"Stale text remains: {stale!r}")

    required = [
        "2022 through 2024",
        "tax year 2022",
        "gross receipts of $50,000 or less",
        "IRS EO BMF",
        "National Center for Charitable Statistics (NCCS) Business Master File (BMF)",
        "18 organizations (about 7%)",
        "10 eligible Black Hills organizations (about 4%)",
        "did not materially change the direction or overall interpretation",
        "local organizations may be accessing some funding channels but receiving smaller amounts",
        "How to read the comparison tables:",
        NEW_MEDIAN_GAP_HEADER,
        "It is not a difference of means.",
        "10,000 bootstrap resamples",
        "Appendix Tables A-2 through A-11",
        "Table A-1 lists the corresponding regional totals",
        "Figures A-1 through A-10",
    ]
    for phrase in required:
        if phrase not in full:
            issues.append(f"Required phrase missing: {phrase!r}")

    data_sources = section_text(
        paragraphs,
        "Data sources",
        "Selection and demographic analysis of benchmark regions",
    ).lower()
    if "analysis used here" in data_sources:
        issues.append("Data-source section still says 'analysis used here'")

    results = section_text(
        paragraphs,
        "Non-profit revenue source comparison (2022)",
        "Number of Non-profit Organizations from 2022 to 2024",
    )
    if "A compact table under each chart lists the median reported dollars by region" not in results:
        issues.append("Narrative does not describe the compact chart-level median tables")
    for caption in MAIN_FIGURES:
        if caption not in results:
            issues.append(f"Main figure missing from narrative block: {caption}")
    for caption in APPENDIX_ONLY_FIGURES:
        if caption in results:
            issues.append(f"Appendix-only figure appears in narrative block: {caption}")
    for prohibited in ["p-value", "p =", "p <", "confidence interval", "95% CI"]:
        if prohibited.lower() in results.lower():
            issues.append(f"Main narrative contains inferential detail: {prohibited!r}")
    main_compact_tables = []
    in_results = False
    for item in body_items(document):
        if isinstance(item, Paragraph) and item.text.strip() == "Non-profit revenue source comparison (2022)":
            in_results = True
            continue
        if isinstance(item, Paragraph) and item.text.strip() == "Number of Non-profit Organizations from 2022 to 2024":
            break
        if in_results and isinstance(item, Table):
            if table_header(item) == ["Region", "Median reported dollars"]:
                main_compact_tables.append(item)
    if len(main_compact_tables) != 5:
        issues.append(f"Main compact median tables: {len(main_compact_tables)} (expect 5)")
    for table in main_compact_tables:
        regions = [row.cells[0].text.strip() for row in table.rows[1:]]
        if regions != REGION_ORDER:
            issues.append(f"Unexpected region order in compact median table: {regions}")

    appendix = section_text(
        paragraphs,
        "Appendix A: Non-profit revenue source comparison details",
        "Table 5-13. Incidence Risk Ratios (IRR) and 95% confidence intervals from Poisson regression, with number of organizations as the response variable, tax year as explanatory variable (reference level 2022), and population count as offset for all organizations.",
    )
    for caption in APPENDIX_FIGURES:
        if caption not in appendix:
            issues.append(f"Figure missing from complete Appendix A chart set: {caption}")
    for caption in SOURCE_TABLE_CAPTIONS:
        if caption not in appendix:
            issues.append(f"Source detail table caption missing from Appendix A: {caption}")
    for stale_caption in [*MAIN_FIGURES, *APPENDIX_ONLY_FIGURES]:
        if stale_caption in appendix:
            issues.append(f"Section-style figure caption remains in Appendix A: {stale_caption}")
    for stale_caption in [*[f"Table 5-{index}:" for index in range(2, 14)]]:
        if stale_caption in appendix:
            issues.append(f"Section-style table caption remains in Appendix A: {stale_caption}")

    for index in range(2, 12):
        caption = next(
            paragraph.text.strip()
            for paragraph in paragraphs
            if paragraph.text.strip().startswith(f"Table A-{index}:")
        )
        table = table_after_caption(document, caption)
        if table is None:
            issues.append(f"No table after {caption}")
            continue
        if table_header(table) != [
            "Region",
            "Reporting rate",
            "Positive reporters (n)",
            "Median reported dollars",
            NEW_MEDIAN_GAP_HEADER,
            "Pairwise p-value vs Black Hills",
        ]:
            issues.append(f"Unexpected header for {caption}: {table_header(table)}")
        regions = [row.cells[0].text.strip() for row in table.rows[1:]]
        if regions != REGION_ORDER:
            issues.append(f"Unexpected region order for {caption}: {regions}")
        if table.rows[1].cells[4].text.strip() != "NA" or table.rows[1].cells[5].text.strip() != "NA":
            issues.append(f"Black Hills comparison cells are not NA in {caption}")

    for table in document.tables:
        header = table_header(table)
        if header and header[0] in {"Revenue source", "Region"}:
            fills = []
            for cell in table.rows[0].cells:
                shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                fills.append(shading.get(qn("w:fill")) if shading is not None else None)
            if any(fill != ADDED_TABLE_HEADER_FILL for fill in fills):
                issues.append(f"Added table header is not light blue: {fills}")
        if not is_q9_added_table(table):
            continue
        if table.style is None or table.style.name != "Table Grid":
            issues.append(f"Q9 table does not use Table Grid style: {header}")
        if table.alignment != WD_TABLE_ALIGNMENT.LEFT:
            issues.append(f"Q9 table is not left-aligned: {header}")
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    expected_alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                        if column_index == 0
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    if paragraph.text.strip() and paragraph.alignment != expected_alignment:
                        issues.append(
                            f"Unexpected Q9 table cell alignment at row {row_index}, "
                            f"column {column_index}: {header}"
                        )
                    for run in paragraph.runs:
                        if not run.text:
                            continue
                        if run.font.name != TABLE_FONT_NAME:
                            issues.append(f"Unexpected Q9 table font {run.font.name!r}: {run.text[:40]!r}")
                        size = run.font.size.pt if run.font.size is not None else None
                        if size != TABLE_FONT_SIZE_PT:
                            issues.append(f"Unexpected Q9 table font size {size!r}: {run.text[:40]!r}")
                        expected_bold = row_index == 0
                        if bool(run.bold) != expected_bold:
                            issues.append(f"Unexpected Q9 table bold state: {run.text[:40]!r}")
                        if run.underline:
                            issues.append(f"Underlined Q9 table text: {run.text[:40]!r}")

    for text in [
        "IRS Exempt Organizations Business Master File (IRS EO BMF)",
        "GivingTuesday 990 DataMart",
        "National Center for Charitable Statistics (NCCS) Business Master File (BMF)",
        "Non-profit revenue analysis",
        "Non-profit revenue source comparison (2022)",
        "Appendix A: Non-profit revenue source comparison details",
    ]:
        check_color(document, text, BLACK, issues)
        heading_matches = [paragraph for paragraph in paragraphs if paragraph.text.strip() == text]
        if heading_matches and heading_matches[0].style.name not in {"Heading 2", "Heading 3"}:
            issues.append(f"Our heading has unexpected style {heading_matches[0].style.name!r}: {text}")
    for text in [
        "Bureau of Labor Statistics Non-profit Research Dataset",
        "Number of Non-profit Organizations from 2022 to 2024",
        "Conclusions",
    ]:
        check_color(document, text, GRAY, issues)

    with ZipFile(DOCX) as zip_file:
        if "word/footnotes.xml" in zip_file.namelist():
            root = etree.fromstring(zip_file.read("word/footnotes.xml"))
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for prefix in [
                "State tax provisions reflect rules in effect for the 2022 tax year",
                "Both r and rs are correlation coefficients",
            ]:
                matches = []
                for footnote in root.xpath("//w:footnote", namespaces=namespace):
                    text = "".join(footnote.xpath(".//w:t/text()", namespaces=namespace)).strip()
                    if text.startswith(prefix):
                        matches.append(footnote)
                if matches:
                    for run in matches[0].xpath(".//w:r", namespaces=namespace):
                        text = "".join(run.xpath(".//w:t/text()", namespaces=namespace))
                        if not text:
                            continue
                        colors = run.xpath("./w:rPr/w:color/@w:val", namespaces=namespace)
                        if colors != [GRAY]:
                            issues.append(f"Teammate footnote is not gray: {prefix!r}")
                            break

    print(f"Auditing: {DOCX.name}")
    print(f"Issues: {len(issues)}")
    for issue in issues:
        print(f"  - {issue}")
    if issues:
        sys.exit(1)
    print("  - No automated issues found")


if __name__ == "__main__":
    main()
