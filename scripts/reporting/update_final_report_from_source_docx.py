"""Color-code the uploaded final report DOCX and regenerate its Markdown copy.

The uploaded DOCX is treated as the current teammate-authored base. This script
colors the base gray, then colors the SWB-added GivingTuesday/Q9 revenue-source
content black. It preserves the uploaded DOCX's tables, charts, and images.
"""

from __future__ import annotations

import html
import re
from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


REPO_ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = REPO_ROOT / "docs" / "deliverables" / "final_report"
SOURCE_DOCX = REPORT_DIR / "source" / "black_hills_philanthropy_source.docx"
WORKING_DIR = REPORT_DIR / "working"
COLOR_DOCX = WORKING_DIR / "black_hills_philanthropy_color_coded.docx"
MD_PATH = WORKING_DIR / "black_hills_philanthropy_color_coded.md"
ASSET_DIR = WORKING_DIR / "assets"

GRAY = "666666"
BLACK = "000000"
ADDED_TABLE_HEADER_FILL = "DEEAF0"
TABLE_FONT_NAME = "Arial"
TABLE_FONT_SIZE_PT = 8
FIGURE_TITLE_FONT_SIZE_PT = 10

IRS_BMF_HEADING = "IRS Exempt Organizations Business Master File (IRS EO BMF)"
NCCS_BMF_HEADING = "National Center for Charitable Statistics (NCCS) Business Master File (BMF)"

IRS_BMF_PARAGRAPHS = [
    (
        "The IRS Exempt Organizations Business Master File (IRS EO BMF) is a public IRS registry "
        "of tax-exempt organizations. Years included: 2022 through 2024 files for "
        "organizations located in the Black Hills and benchmark-region counties."
    ),
    (
        "The file is useful for identifying organizations, locations, exemption types, ruling "
        "dates, and broad classification fields. In this report it supports counts of non-profit "
        "organizations, geography checks, and classification details that are not always complete "
        "in the GivingTuesday files."
    ),
    (
        "The IRS EO BMF is a registry file, not a detailed financial filing file. It does not "
        "include the full revenue, expense, contribution, grant, or asset detail used in the "
        "detailed revenue-source tables."
    ),
    "The IRS EO BMF appendix provides additional details on source files, fields, and processing steps.",
]

GIVINGTUESDAY_PARAGRAPHS = [
    (
        "The GivingTuesday 990 DataMart is the main source used for non-profit financial "
        "information in this report. It is based on public IRS filings submitted by tax-exempt "
        "organizations, including Form 990, Form 990-EZ, and Form 990-PF. Years included: "
        "selected GivingTuesday files for tax years 2022 through 2024. The detailed "
        "revenue-source tables use tax year 2022."
    ),
    (
        "The IRS form type matters because different organizations file different versions of "
        "the Form 990 family. Form 990 is the full annual return for larger non-profit "
        "organizations and provides the most detail. Form 990-EZ is a shorter return used by "
        "some smaller organizations and includes less detail. Form 990-PF is used by private "
        "foundations. Form 990-N, also called the e-Postcard, is generally used by very small "
        "organizations with gross receipts of $50,000 or less. Form 990-N contains only limited "
        "information, so those organizations are not included in the detailed GivingTuesday "
        "financial tables used here."
    ),
    (
        "The project uses the GivingTuesday Basic Fields files for Form 990, Form 990-EZ, and "
        "Form 990-PF, limited to organizations located in the Black Hills and benchmark-region "
        "counties. When GivingTuesday did not include all classification details needed to group "
        "organizations consistently, the project supplemented those details with the National "
        "Center for Charitable Statistics (NCCS) Business Master File (BMF) and the IRS EO BMF."
    ),
    (
        "The main strength of the GivingTuesday source is that it provides more financial detail "
        "than registry-only sources. This makes it useful for describing organization size, "
        "revenue sources, and overall financial composition. However, it should not be read as a "
        "complete count of every non-profit or every charitable activity in a region."
    ),
    (
        "The source also has limits because the IRS forms do not all ask for the same level of "
        "detail. Organizations filing the full Form 990 report more detail about specific "
        "revenue sources than organizations filing Form 990-EZ or Form 990-PF. Very small "
        "organizations that file only Form 990-N, organizations not required to file public "
        "Form 990-family returns, houses of worship and some related organizations, and informal "
        "giving are not fully represented."
    ),
    (
        "For this reason, the report uses GivingTuesday as a practical view of reported "
        "non-profit finances on IRS filings. It does not capture all generosity in the region, "
        "and it should be interpreted alongside local knowledge and the registry sources "
        "described above."
    ),
    "The appendices provide additional detail on the source files and fields used.",
]

NCCS_BMF_PARAGRAPHS = [
    (
        "The National Center for Charitable Statistics (NCCS) Business Master File (BMF) is a "
        "registry-style non-profit dataset maintained by the Urban Institute using IRS "
        "exempt-organization records. Years included: 2022 through 2024 NCCS BMF files "
        "for organizations in the Black Hills and benchmark regions."
    ),
    (
        "The NCCS BMF is useful for organization classification, including National Taxonomy of "
        "Exempt Entities (NTEE) activity codes, and for checking organization counts and "
        "geography over time."
    ),
    (
        "Like the IRS EO BMF, the NCCS BMF is not the main source for detailed revenue-source "
        "fields. It does not provide the full Form 990 line detail needed to compare "
        "contributions, government grants, membership dues, fundraising events, and other "
        "revenue sources."
    ),
    "The NCCS BMF appendix provides additional details on source files, fields, and processing steps.",
]

Q9_RELEVANT_DETAILS_ITEMS = [
    ("Heading 3", "Non-profit revenue analysis"),
    ("Normal", "This section uses the GivingTuesday Form 990-family data described under Data sources."),
    (
        "Normal",
        (
            "For tax year 2022, after requiring positive total revenue, the primary analysis file "
            "includes 1,799 filing records for 1,797 organizations: 422 in the Black Hills, 332 "
            "in Billings, 197 in Flagstaff, 283 in Missoula, and 565 in Sioux Falls."
        ),
    ),
    (
        "Normal",
        (
            "We also checked the 2022 revenue-source comparisons with and without 25 records for "
            "hospitals, universities, and political organizations. Including those records did "
            "not materially change the direction or overall interpretation of the patterns. The "
            "revenue-source results presented below use the version without those organizations "
            "so the comparison focuses on more similar non-profit peers."
        ),
    ),
    ("Normal", "Definitions for the revenue sources analyzed in this section are provided in Appendix Table A-12."),
    (
        "Normal",
        (
            "The data cannot cleanly separate individual giving from foundation grants, corporate "
            "gifts, bequests, or other institutional support. That limitation applies especially "
            "to total contributions and mixed or unclassified contributions. These categories "
            "describe what organizations reported on IRS filings, not who ultimately made each gift."
        ),
    ),
    (
        "Normal",
        (
            "This report presents two complementary views of the same filing data. Regional totals "
            "(Appendix Table A-1) sum reported dollars across all organizations in the analysis "
            "file for each source and region. That is the direct read on how much money flowed "
            "through each channel in the overall landscape, but a few very large filers can "
            "dominate those sums, so totals can differ sharply from what most organizations experience."
        ),
    ),
    (
        "Normal",
        (
            "The headline comparisons in this section use the median reported dollar amount among "
            "organizations that reported a positive amount for that source. That answers a "
            "different question: when a local organization does use the channel, how large is the "
            "amount it reports? The median is less pulled up or down by a handful of extreme "
            "filers than a regional total or average would be, and it compares only organizations "
            "that actually reported the source rather than treating every non-reporter as zero. "
            "That separation matters because reporting rates differ by source and region, and on "
            "IRS forms a missing line does not always mean the organization earned nothing from "
            "that source. Reporting rates and reporter counts appear in Appendix Tables A-2 "
            "through A-11; regional totals appear in Appendix Table A-1."
        ),
    ),
    ("Heading 3", "How many organizations could we compare?"),
    (
        "Normal",
        (
            "Only organizations that reported income from a given revenue source were included in "
            "that source's comparison. The number of organizations varies widely by source because "
            "different revenue lines appear on different IRS forms and because many organizations "
            "do not use every funding channel."
        ),
    ),
    (
        "Normal",
        (
            "The 422 Black Hills records above are all Form 990-family filers with positive total "
            "revenue in 2022. Comparisons for Form 990 Part VIII contribution lines (government "
            "grants, federated campaigns, related-organization contributions, membership dues, "
            "and fundraising events) use only organizations that filed the full Form 990. That "
            "leaves 256 eligible Black Hills records for those detailed contribution lines. "
            "Organizations filing Form 990-EZ or Form 990-PF are not counted as zero for those "
            "lines because their forms do not provide the same line-level detail."
        ),
    ),
    (
        "Normal",
        (
            "For total revenue, total contributions, program service revenue, mixed or "
            "unclassified contributions, and other revenue, most pairwise comparisons involved "
            "dozens or hundreds of positive reporters in each region."
        ),
    ),
    (
        "Normal",
        (
            "Some Part VIII lines are reported by far fewer organizations. Among the 256 eligible "
            "Black Hills Form 990 records, 18 organizations (about 7%) reported positive "
            "federated campaign contributions. Some benchmark comparisons for that source "
            "involved as few as three to five organizations with positive amounts. "
            "Related-organization contributions were also uncommon: 10 eligible Black Hills "
            "organizations (about 4%) reported a positive amount. Fundraising event contributions "
            "and membership dues fell in between, with moderate sample sizes in some benchmark pairs."
        ),
    ),
    (
        "Normal",
        (
            "Patterns for sources with very small reporter groups should be read with extra "
            "caution because a few organizations can have a large effect on the median. The "
            "reporting rates in Appendix Tables A-2 through A-11 show how common each revenue "
            "source was in each region."
        ),
    ),
]

SELECTED_MAIN_BODY_FIGURES = [
    "Figure 5-2: Total Revenue by Region",
    "Figure 5-3: Program Service Revenue by Region",
    "Figure 5-4: Total Contributions by Region",
    "Figure 5-5: Government Grants Received by Region",
    "Figure 5-9: Fundraising Event Contributions by Region",
]

APPENDIX_ONLY_FIGURES = [
    "Figure 5-6: Federated Campaign Contributions by Region",
    "Figure 5-7: Related Organization Contributions by Region",
    "Figure 5-8: Membership Dues by Region",
    "Figure 5-10: Mixed or Unclassified Contributions by Region",
    "Figure 5-11: Other Revenue by Region",
]

APPENDIX_FIGURE_TITLE_MAP = {
    "Figure 5-2: Total Revenue by Region": "Figure A-1: Total Revenue by Region",
    "Figure 5-3: Program Service Revenue by Region": "Figure A-2: Program Service Revenue by Region",
    "Figure 5-4: Total Contributions by Region": "Figure A-3: Total Contributions by Region",
    "Figure 5-5: Government Grants Received by Region": "Figure A-4: Government Grants Received by Region",
    "Figure 5-6: Federated Campaign Contributions by Region": "Figure A-5: Federated Campaign Contributions by Region",
    "Figure 5-7: Related Organization Contributions by Region": "Figure A-6: Related Organization Contributions by Region",
    "Figure 5-8: Membership Dues by Region": "Figure A-7: Membership Dues by Region",
    "Figure 5-9: Fundraising Event Contributions by Region": "Figure A-8: Fundraising Event Contributions by Region",
    "Figure 5-10: Mixed or Unclassified Contributions by Region": "Figure A-9: Mixed or Unclassified Contributions by Region",
    "Figure 5-11: Other Revenue by Region": "Figure A-10: Other Revenue by Region",
}

APPENDIX_TABLE_TITLE_MAP = {
    "Table 5-13: Total Reported Dollars by Revenue Source and Region (Aggregate)": "Table A-1: Total Reported Dollars by Revenue Source and Region (Aggregate)",
    "Table 5-2: Total Revenue by Region": "Table A-2: Total Revenue by Region",
    "Table 5-3: Program Service Revenue by Region": "Table A-3: Program Service Revenue by Region",
    "Table 5-4: Total Contributions by Region": "Table A-4: Total Contributions by Region",
    "Table 5-5: Government Grants Received by Region": "Table A-5: Government Grants Received by Region",
    "Table 5-6: Federated Campaign Contributions by Region": "Table A-6: Federated Campaign Contributions by Region",
    "Table 5-7: Related Organization Contributions by Region": "Table A-7: Related Organization Contributions by Region",
    "Table 5-8: Membership Dues by Region": "Table A-8: Membership Dues by Region",
    "Table 5-9: Fundraising Event Contributions by Region": "Table A-9: Fundraising Event Contributions by Region",
    "Table 5-10: Mixed or Unclassified Contributions by Region": "Table A-10: Mixed or Unclassified Contributions by Region",
    "Table 5-11: Other Revenue by Region": "Table A-11: Other Revenue by Region",
    "Table 5-12: Revenue Source Definitions": "Table A-12: Revenue Source Definitions",
}

APPENDIX_ONLY_FIGURE_ASSET_MAP = {
    "Figure 5-6: Federated Campaign Contributions by Region": REPO_ROOT
    / "docs"
    / "analysis"
    / "revenue_sources"
    / "assets"
    / "q9_2022"
    / "client_2022_pairwise_positive_median_federated_campaigns.png",
    "Figure 5-7: Related Organization Contributions by Region": REPO_ROOT
    / "docs"
    / "analysis"
    / "revenue_sources"
    / "assets"
    / "q9_2022"
    / "client_2022_pairwise_positive_median_related_org_contributions.png",
    "Figure 5-8: Membership Dues by Region": REPO_ROOT
    / "docs"
    / "analysis"
    / "revenue_sources"
    / "assets"
    / "q9_2022"
    / "client_2022_pairwise_positive_median_membership_dues.png",
    "Figure 5-10: Mixed or Unclassified Contributions by Region": REPO_ROOT
    / "docs"
    / "analysis"
    / "revenue_sources"
    / "assets"
    / "q9_2022"
    / "client_2022_pairwise_positive_median_mixed_unclassified_contributions.png",
    "Figure 5-11: Other Revenue by Region": REPO_ROOT
    / "docs"
    / "analysis"
    / "revenue_sources"
    / "assets"
    / "q9_2022"
    / "client_2022_pairwise_positive_median_residual_other_revenue.png",
}

Q9_FIGURE_TITLES = SELECTED_MAIN_BODY_FIGURES + list(APPENDIX_FIGURE_TITLE_MAP.values())
Q9_TABLE_TITLES = [
    "Table 5-2a: Median Reported Dollars for Total Revenue by Region",
    "Table 5-3a: Median Reported Dollars for Program Service Revenue by Region",
    "Table 5-4a: Median Reported Dollars for Total Contributions by Region",
    "Table 5-5a: Median Reported Dollars for Government Grants Received by Region",
    "Table 5-9a: Median Reported Dollars for Fundraising Event Contributions by Region",
    *APPENDIX_TABLE_TITLE_MAP.values(),
]

OLD_MEDIAN_GAP_HEADER = "Black Hills minus benchmark region (95% CI)"
NEW_MEDIAN_GAP_HEADER = "Black Hills median minus benchmark median (95% bootstrap CI)"

APPENDIX_INTRO = (
    "The appendix provides the complete 2022 revenue-source comparison details. Tables A-2 "
    "through A-11 provide reporting rates, reporter counts, medians, median gaps, confidence "
    "intervals, and p-values for each revenue source. Table A-12 defines the revenue sources. "
    "Table A-1 lists the corresponding regional totals summed across all organizations in the "
    "analysis file. Figures A-1 through A-10 show the complete source-by-source chart set, "
    "including the charts also shown in the main narrative. The comparison columns apply only "
    "to each benchmark region versus Black Hills; no Black Hills-only p-value or median-gap "
    "confidence interval exists, so those Black Hills cells are marked NA. Rare revenue sources "
    "have fewer reporting organizations; see the sample-size discussion in the non-profit "
    "revenue analysis section."
)

AGGREGATE_TOTALS_NOTE = (
    "Table A-1 sums reported dollars for each revenue source and region over that source's "
    "eligible reporting universe in the 2022 analysis file. Totals for total revenue, total "
    "contributions, mixed or unclassified contributions, and other revenue include all "
    "organizations in the file. Government grants, federated campaigns, related-organization "
    "contributions, membership dues, and fundraising event contributions include Form 990 filers "
    "only; 990-EZ and 990-PF filers do not report those Part VIII sub-lines and are excluded "
    "from those rows, not counted as zero. Within Form 990, a blank sub-line is treated as zero. "
    "Program service revenue includes organizations for which that line is reported on the filed "
    "form. Regional totals reflect both universe size and report size; a few very large filers "
    "can dominate the sum. The compact tables under the main-text charts report medians among "
    "positive reporters only, which is the basis for the pairwise comparisons in Tables A-2 "
    "through A-11."
)

COMPARISON_TABLE_METHOD_NOTE = (
    "How to read the comparison tables: In Tables A-2 through A-11, the column "
    f"labeled \"{NEW_MEDIAN_GAP_HEADER}\" compares the median reported dollars "
    "among organizations that reported a positive amount for that revenue source. "
    "It is not a difference of means. A negative value means the Black Hills median "
    "was lower than the benchmark-region median; a positive value means it was "
    "higher. The 95% confidence interval was calculated using 10,000 bootstrap "
    "resamples within each region and reflects uncertainty around the median difference."
)


def iter_body_items(document: Document):
    for element in document.element.body.iterchildren():
        if element.tag.endswith("}p"):
            yield Paragraph(element, document._body)
        elif element.tag.endswith("}tbl"):
            yield Table(element, document._body)


def body_children(document: Document) -> list:
    return list(document.element.body.iterchildren())


def is_paragraph_element(element) -> bool:
    return element.tag.endswith("}p")


def is_table_element(element) -> bool:
    return element.tag.endswith("}tbl")


def element_text(element) -> str:
    if not is_paragraph_element(element):
        return ""
    return "".join(text.text or "" for text in element.findall(".//" + qn("w:t"))).strip()


def element_has_image(element) -> bool:
    return is_paragraph_element(element) and bool(element.findall(".//" + qn("a:blip")))


def paragraph_from_element(document: Document, element) -> Paragraph:
    return Paragraph(element, document._body)


def clear_paragraph_runs(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    clear_paragraph_runs(paragraph)
    if text:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")


def find_paragraph_index(
    document: Document,
    candidates: str | list[str],
    *,
    start: int = 0,
) -> int:
    if isinstance(candidates, str):
        candidate_set = {candidates}
    else:
        candidate_set = set(candidates)
    for index, element in enumerate(body_children(document)[start:], start=start):
        if is_paragraph_element(element) and element_text(element) in candidate_set:
            return index
    raise ValueError(f"Could not find paragraph: {sorted(candidate_set)}")


def make_paragraph_element(document: Document, text: str, style: str = "Normal"):
    paragraph = document.add_paragraph()
    paragraph.style = style
    paragraph.add_run(text)
    element = paragraph._element
    element.getparent().remove(element)
    return element


def make_paragraph_elements(document: Document, items: list[tuple[str, str]]) -> list:
    return [make_paragraph_element(document, text, style) for style, text in items]


def make_image_paragraph_element(document: Document, image_path: Path):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
    element = paragraph._element
    element.getparent().remove(element)
    return element


def insert_elements_before(document: Document, reference, elements: list) -> None:
    body = document.element.body
    for element in elements:
        body.insert(body.index(reference), element)


def remove_elements(document: Document, elements: list) -> None:
    body = document.element.body
    for element in elements:
        if element.getparent() is body:
            body.remove(element)


def replace_heading_text(
    document: Document,
    candidates: str | list[str],
    new_text: str,
) -> None:
    index = find_paragraph_index(document, candidates)
    set_paragraph_text(paragraph_from_element(document, body_children(document)[index]), new_text)


def replace_content_between_headings(
    document: Document,
    start_heading_candidates: str | list[str],
    end_heading_candidates: str | list[str],
    paragraph_texts: list[str],
) -> None:
    start_index = find_paragraph_index(document, start_heading_candidates)
    end_index = find_paragraph_index(document, end_heading_candidates, start=start_index + 1)
    children = body_children(document)
    reference = children[end_index]
    remove_elements(document, children[start_index + 1 : end_index])
    elements = [make_paragraph_element(document, text) for text in paragraph_texts]
    insert_elements_before(document, reference, elements)


def replace_block(
    document: Document,
    start_heading: str,
    end_heading: str,
    elements: list,
) -> None:
    start_index = find_paragraph_index(document, start_heading)
    end_index = find_paragraph_index(document, end_heading, start=start_index + 1)
    children = body_children(document)
    reference = children[end_index]
    remove_elements(document, children[start_index:end_index])
    insert_elements_before(document, reference, elements)


def clone_caption_and_next_table(document: Document, caption_text: str) -> list:
    caption_index = find_paragraph_index(document, caption_text)
    children = body_children(document)
    for index in range(caption_index + 1, len(children)):
        if is_table_element(children[index]):
            return [deepcopy(children[caption_index]), deepcopy(children[index])]
        if is_paragraph_element(children[index]) and element_text(children[index]):
            break
    raise ValueError(f"Could not find table after caption: {caption_text}")


def clone_caption_and_next_image(document: Document, caption_text: str) -> list:
    caption_index = find_paragraph_index(document, caption_text)
    children = body_children(document)
    for index in range(caption_index + 1, len(children)):
        if element_has_image(children[index]):
            return [deepcopy(children[caption_index]), deepcopy(children[index])]
        if is_paragraph_element(children[index]) and element_text(children[index]):
            break
        if is_table_element(children[index]):
            break
    raise ValueError(f"Could not find image after caption: {caption_text}")


def clone_caption_and_next_image_after(
    document: Document,
    caption_text: str | list[str],
    *,
    after_heading: str,
) -> list:
    start = find_paragraph_index(document, after_heading)
    caption_index = find_paragraph_index(document, caption_text, start=start + 1)
    children = body_children(document)
    for index in range(caption_index + 1, len(children)):
        if element_has_image(children[index]):
            return [deepcopy(children[caption_index]), deepcopy(children[index])]
        if is_paragraph_element(children[index]) and element_text(children[index]):
            break
        if is_table_element(children[index]):
            break
    raise ValueError(f"Could not find image after caption: {caption_text}")


def table_after_caption(document: Document, caption_text: str | list[str]) -> Table:
    caption_index = find_paragraph_index(document, caption_text)
    children = body_children(document)
    for index in range(caption_index + 1, len(children)):
        if is_table_element(children[index]):
            return Table(children[index], document._body)
        if is_paragraph_element(children[index]) and element_text(children[index]):
            break
    raise ValueError(f"Could not find table after caption: {caption_text}")


def make_table_element(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text)
    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            set_cell_text(cell, text)
    element = table._element
    element.getparent().remove(element)
    return element


def compact_median_table_from_detail(document: Document, detail_caption: str):
    source = table_after_caption(document, detail_caption)
    headers = [cell.text.strip() for cell in source.rows[0].cells]
    region_index = headers.index("Region")
    median_index = headers.index("Median reported dollars")
    rows = [
        [
            row.cells[region_index].text.strip(),
            row.cells[median_index].text.strip(),
        ]
        for row in source.rows[1:]
    ]
    return make_table_element(document, ["Region", "Median reported dollars"], rows)


def remove_caption_and_next_table_after(
    document: Document,
    caption_text: str | list[str],
    after_heading: str,
) -> None:
    try:
        start = find_paragraph_index(document, after_heading)
        caption_index = find_paragraph_index(document, caption_text, start=start + 1)
    except ValueError:
        return
    children = body_children(document)
    to_remove = [children[caption_index]]
    for index in range(caption_index + 1, len(children)):
        if is_table_element(children[index]):
            to_remove.append(children[index])
            break
        if is_paragraph_element(children[index]) and element_text(children[index]):
            break
    remove_elements(document, to_remove)


def remove_caption_and_next_image_after(
    document: Document,
    caption_text: str | list[str],
    after_heading: str,
) -> None:
    try:
        start = find_paragraph_index(document, after_heading)
        caption_index = find_paragraph_index(document, caption_text, start=start + 1)
    except ValueError:
        return
    children = body_children(document)
    to_remove = [children[caption_index]]
    for index in range(caption_index + 1, len(children)):
        if element_has_image(children[index]):
            to_remove.append(children[index])
            break
        if is_paragraph_element(children[index]) and element_text(children[index]):
            break
        if is_table_element(children[index]):
            break
    remove_elements(document, to_remove)


def insert_elements_before_caption(
    document: Document,
    target_caption: str | list[str],
    elements: list,
    *,
    after_heading: str,
) -> None:
    start = find_paragraph_index(document, after_heading)
    target_index = find_paragraph_index(document, target_caption, start=start + 1)
    reference = body_children(document)[target_index]
    insert_elements_before(document, reference, [deepcopy(element) for element in elements])


def image_elements_with_new_caption(
    document: Document,
    caption_elements: list,
    new_caption: str,
) -> list:
    return [
        make_paragraph_element(document, new_caption, "Table/Figure Title"),
        *[deepcopy(element) for element in caption_elements if element_has_image(element)],
    ]


def retitle_paragraphs_after_heading(
    document: Document,
    title_map: dict[str, str],
    *,
    after_heading: str,
) -> None:
    try:
        start = find_paragraph_index(document, after_heading)
    except ValueError:
        return
    children = body_children(document)
    for index in range(start + 1, len(children)):
        element = children[index]
        if not is_paragraph_element(element):
            continue
        text = element_text(element).strip()
        if text in title_map:
            paragraph = paragraph_from_element(document, element)
            set_paragraph_text(paragraph, title_map[text])


def remove_paragraph_starting_with_after(
    document: Document,
    prefix: str,
    *,
    after_heading: str,
) -> None:
    try:
        start = find_paragraph_index(document, after_heading)
    except ValueError:
        return
    children = body_children(document)
    to_remove = [
        element
        for element in children[start + 1 :]
        if is_paragraph_element(element) and element_text(element).startswith(prefix)
    ]
    remove_elements(document, to_remove)


def replace_first_paragraph_after_heading(
    document: Document,
    heading_text: str,
    new_text: str,
) -> None:
    heading_index = find_paragraph_index(document, heading_text)
    children = body_children(document)
    for index in range(heading_index + 1, len(children)):
        if is_paragraph_element(children[index]):
            paragraph = paragraph_from_element(document, children[index])
            if paragraph.text.strip():
                set_paragraph_text(paragraph, new_text)
                return
    raise ValueError(f"Could not find paragraph after heading: {heading_text}")


def replace_paragraph_starting_with(document: Document, prefix: str | list[str], new_text: str) -> None:
    prefixes = [prefix] if isinstance(prefix, str) else prefix
    for element in body_children(document):
        if not is_paragraph_element(element):
            continue
        if any(element_text(element).startswith(candidate) for candidate in prefixes):
            set_paragraph_text(paragraph_from_element(document, element), new_text)
            return
    raise ValueError(f"Could not find paragraph starting with: {prefixes}")


def apply_q9_report_styles(document: Document) -> None:
    title_set = set(Q9_FIGURE_TITLES + Q9_TABLE_TITLES)
    figure_title_set = set(Q9_FIGURE_TITLES)
    has_table_figure_title = "Table/Figure Title" in [style.name for style in document.styles]
    for paragraph in document.paragraphs:
        title_text = paragraph.text.strip()
        if title_text not in title_set:
            continue
        if has_table_figure_title:
            paragraph.style = "Table/Figure Title"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            if run.text:
                run.font.name = TABLE_FONT_NAME
                if title_text in figure_title_set:
                    run.font.size = Pt(FIGURE_TITLE_FONT_SIZE_PT)
                run.bold = True
                run.underline = False


def update_q9_source_table_headers(document: Document) -> None:
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if OLD_MEDIAN_GAP_HEADER not in header:
            continue
        column_index = header.index(OLD_MEDIAN_GAP_HEADER)
        set_cell_text(table.rows[0].cells[column_index], NEW_MEDIAN_GAP_HEADER)


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("a:blip")))


def set_run_color(run_element, color: str) -> None:
    run_properties = run_element.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        run_element.insert(0, run_properties)
    color_element = run_properties.find(qn("w:color"))
    if color_element is None:
        color_element = OxmlElement("w:color")
        run_properties.append(color_element)
    color_element.set(qn("w:val"), color)


def set_paragraph_color(paragraph: Paragraph, color: str) -> None:
    for run in paragraph._element.findall(".//" + qn("w:r")):
        set_run_color(run, color)


def set_table_color(table: Table, color: str) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_color(paragraph, color)


def shade_header(table: Table) -> None:
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), ADDED_TABLE_HEADER_FILL)


def is_q9_added_table(table: Table) -> bool:
    header = tuple(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ()
    return header in {
        ("Region", "Median reported dollars"),
        (
            "Region",
            "Reporting rate",
            "Positive reporters (n)",
            "Median reported dollars",
            NEW_MEDIAN_GAP_HEADER,
            "Pairwise p-value vs Black Hills",
        ),
        ("Revenue source", "Black Hills", "Billings", "Flagstaff", "Missoula", "Sioux Falls"),
        ("Revenue source", "What it measures", "How it is reported on IRS filings"),
    }


def format_q9_table(table: Table) -> None:
    if not table.rows:
        return
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    shade_header(table)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    if not run.text:
                        continue
                    run.font.name = TABLE_FONT_NAME
                    run.font.size = Pt(TABLE_FONT_SIZE_PT)
                    run.bold = row_index == 0
                    run.underline = False


def format_q9_tables(document: Document) -> None:
    for table in document.tables:
        if is_q9_added_table(table):
            format_q9_table(table)


def build_q9_results_elements(
    document: Document,
    selected_figures: dict[str, list],
    compact_median_tables: dict[str, object],
) -> list:
    def add_figure_with_compact_table(caption: str, table_title: str) -> None:
        elements.extend(deepcopy(element) for element in selected_figures[caption])
        elements.append(make_paragraph_element(document, table_title, "Table/Figure Title"))
        elements.append(deepcopy(compact_median_tables[caption]))

    elements: list = make_paragraph_elements(
        document,
        [
            ("Heading 3", "Non-profit revenue source comparison (2022)"),
            (
                "Normal",
                (
                    "We compared whether non-profit revenue sources differ between the Black Hills "
                    "and the benchmark regions using 2022 filings. The main comparisons use "
                    "median dollars among organizations that reported a positive amount for each "
                    "source. Complete technical details are provided in Appendix A. The pattern "
                    "varies by revenue source and benchmark region."
                ),
            ),
            (
                "Normal",
                (
                    "The figures in this section show the revenue sources that carry the main "
                    "story: total revenue, total contributions, program service revenue, government "
                    "grants, and fundraising events. A compact table under each chart lists the "
                    "median reported dollars by region for that revenue source. Additional "
                    "source-specific charts and all detailed comparison tables appear in Appendix A."
                ),
            ),
        ],
    )
    elements.extend(
        make_paragraph_elements(
            document,
            [
                ("Heading 3", "What the 2022 revenue-source comparison shows"),
                (
                    "Normal",
                    (
                        "Across the ten revenue sources, the Black Hills medians were often near "
                        "the lower end of the regional range, but the pattern was not uniform; the "
                        "size and direction of each gap depended on the source and the benchmark region."
                    ),
                ),
                (
                    "Normal",
                    (
                        "For the two broadest measures, the regions were broadly similar. The Black "
                        "Hills median total revenue was $176,760, somewhat below the benchmark "
                        "medians of about $208,000 to $222,000, though that gap is small next to "
                        "the wide spread in organization size within every region. The Black Hills "
                        "median total contributions was $110,308, lower than Billings, Flagstaff, "
                        "and Missoula but higher than Sioux Falls ($100,000)."
                    ),
                ),
            ],
        )
    )
    add_figure_with_compact_table(
        "Figure 5-2: Total Revenue by Region",
        "Table 5-2a: Median Reported Dollars for Total Revenue by Region",
    )
    add_figure_with_compact_table(
        "Figure 5-4: Total Contributions by Region",
        "Table 5-4a: Median Reported Dollars for Total Contributions by Region",
    )
    elements.extend(
        make_paragraph_elements(
            document,
            [
                (
                    "Normal",
                    (
                        "Several sources showed wider gaps, with the Black Hills median well below "
                        "the benchmark: program service revenue versus Flagstaff ($106,346 versus "
                        "$188,968); government grants versus Billings ($163,367 versus $274,793) "
                        "and Flagstaff ($163,367 versus $468,152); fundraising event contributions "
                        "versus Billings, Missoula, and Sioux Falls ($12,427 versus $38,516, "
                        "$41,756, and $70,761); other revenue versus Billings and Sioux Falls "
                        "($14,805 versus $27,630 and $32,467); mixed or unclassified contributions "
                        "versus Flagstaff ($65,792 versus $92,278); and federated campaign "
                        "contributions versus Sioux Falls ($30,096 versus $150,655). Across these "
                        "sources the gap ran the same way, with the Black Hills amount the lower of "
                        "the two."
                    ),
                ),
            ],
        )
    )
    add_figure_with_compact_table(
        "Figure 5-3: Program Service Revenue by Region",
        "Table 5-3a: Median Reported Dollars for Program Service Revenue by Region",
    )
    elements.extend(
        make_paragraph_elements(
            document,
            [
                (
                    "Normal",
                    (
                        "Participation and dollar amounts can tell different stories. Government "
                        "grants are a useful example: about 44% of eligible Black Hills "
                        "organizations reported this source compared with about 34% in Billings, "
                        "while the median reported amount was $163,367 in Black Hills and $274,793 "
                        "in Billings. A region can have more organizations using a revenue source "
                        "while still showing a lower typical dollar amount among the organizations "
                        "that use it."
                    ),
                ),
            ],
        )
    )
    add_figure_with_compact_table(
        "Figure 5-5: Government Grants Received by Region",
        "Table 5-5a: Median Reported Dollars for Government Grants Received by Region",
    )
    elements.extend(
        make_paragraph_elements(
            document,
            [
                (
                    "Normal",
                    (
                        "Fundraising event contributions are also important for the client-facing "
                        "story because they describe a familiar local fundraising channel. The "
                        "Black Hills median was $12,427, close to Flagstaff ($11,480) but below "
                        "Billings, Missoula, and Sioux Falls."
                    ),
                ),
            ],
        )
    )
    add_figure_with_compact_table(
        "Figure 5-9: Fundraising Event Contributions by Region",
        "Table 5-9a: Median Reported Dollars for Fundraising Event Contributions by Region",
    )
    elements.extend(
        make_paragraph_elements(
            document,
            [
                (
                    "Normal",
                    (
                        "Black Hills was not lower everywhere. Its median was the higher one in "
                        "at least one comparison for total contributions, membership dues, "
                        "federated campaign contributions, and related-organization contributions. "
                        "For membership dues the regional medians were similar in size, and "
                        "related-organization contributions rest on very few reporters (noted above), so that "
                        "ordering is unstable. Complete regional values and reporter counts for "
                        "every comparison appear in Appendix Tables A-2 through A-11."
                    ),
                ),
                (
                    "Normal",
                    (
                        "The connection to the Black Hills Area Community Foundation is most "
                        "direct where participation and dollar amounts point in different "
                        "directions. For example, a relatively large share of eligible Black Hills "
                        "organizations reported government grants, but the typical grant amount "
                        "among those organizations was lower than in every benchmark region. Black "
                        "Hills organizations also reported lower typical amounts for program "
                        "service revenue and several other sources. Together, these "
                        "patterns suggest that local organizations may be accessing some funding "
                        "channels but receiving smaller amounts through them. If this interpretation "
                        "matches local experience, the Foundation could use it to inform "
                        "conversations about revenue diversification, grant-seeking capacity, and "
                        "financial planning. These data do not show why "
                        "the patterns occur or describe the circumstances of any individual organization."
                    ),
                ),
                (
                    "Normal",
                    (
                        "These findings describe observed 2022 revenue patterns, not causal "
                        "explanations or proof that the regions differ in every comparable "
                        "organization. Differences in organization type, local service needs, "
                        "government funding, tourism-related activity, or the presence of large "
                        "institutions could contribute but would require context beyond these "
                        "filings. The results also do not separate individual giving from "
                        "institutional giving. Appendix Tables A-2 through A-11 provide the "
                        "complete numerical and technical comparison details."
                    ),
                ),
            ],
        )
    )
    return elements


def patch_report_content(document: Document) -> None:
    selected_figures = {
        caption: clone_caption_and_next_image(document, caption)
        for caption in SELECTED_MAIN_BODY_FIGURES
    }
    appendix_figure_elements = {
        caption: selected_figures[caption]
        for caption in SELECTED_MAIN_BODY_FIGURES
    }
    for caption in APPENDIX_ONLY_FIGURES:
        try:
            appendix_figure_elements[caption] = clone_caption_and_next_image_after(
                document,
                [caption, APPENDIX_FIGURE_TITLE_MAP[caption]],
                after_heading="Appendix A: Non-profit revenue source comparison details",
            )
        except ValueError:
            appendix_figure_elements[caption] = [
                make_paragraph_element(document, caption, "Table/Figure Title"),
                make_image_paragraph_element(document, APPENDIX_ONLY_FIGURE_ASSET_MAP[caption]),
            ]
    compact_median_tables = {
        "Figure 5-2: Total Revenue by Region": compact_median_table_from_detail(
            document,
            ["Table 5-2: Total Revenue by Region", "Table A-2: Total Revenue by Region"],
        ),
        "Figure 5-3: Program Service Revenue by Region": compact_median_table_from_detail(
            document,
            ["Table 5-3: Program Service Revenue by Region", "Table A-3: Program Service Revenue by Region"],
        ),
        "Figure 5-4: Total Contributions by Region": compact_median_table_from_detail(
            document,
            ["Table 5-4: Total Contributions by Region", "Table A-4: Total Contributions by Region"],
        ),
        "Figure 5-5: Government Grants Received by Region": compact_median_table_from_detail(
            document,
            ["Table 5-5: Government Grants Received by Region", "Table A-5: Government Grants Received by Region"],
        ),
        "Figure 5-9: Fundraising Event Contributions by Region": compact_median_table_from_detail(
            document,
            ["Table 5-9: Fundraising Event Contributions by Region", "Table A-9: Fundraising Event Contributions by Region"],
        ),
    }

    replace_heading_text(
        document,
        ["IRS exempt organizations business master file", IRS_BMF_HEADING],
        IRS_BMF_HEADING,
    )
    replace_heading_text(
        document,
        ["NCCS business master file", NCCS_BMF_HEADING],
        NCCS_BMF_HEADING,
    )
    replace_content_between_headings(
        document,
        IRS_BMF_HEADING,
        "GivingTuesday 990 DataMart",
        IRS_BMF_PARAGRAPHS,
    )
    replace_content_between_headings(
        document,
        "GivingTuesday 990 DataMart",
        NCCS_BMF_HEADING,
        GIVINGTUESDAY_PARAGRAPHS,
    )
    replace_content_between_headings(
        document,
        NCCS_BMF_HEADING,
        "Bureau of Labor Statistics Non-profit Research Dataset",
        NCCS_BMF_PARAGRAPHS,
    )

    replace_block(
        document,
        "Non-profit revenue analysis",
        "Results of non-profit organization analysis",
        make_paragraph_elements(document, Q9_RELEVANT_DETAILS_ITEMS),
    )
    replace_block(
        document,
        "Non-profit revenue source comparison (2022)",
        "Number of Non-profit Organizations from 2022 to 2024",
        build_q9_results_elements(document, selected_figures, compact_median_tables),
    )

    remove_caption_and_next_table_after(
        document,
        "Table 5-1: Median Reported Dollars by Revenue Source and Region",
        "Appendix A: Non-profit revenue source comparison details",
    )
    for caption in APPENDIX_FIGURE_TITLE_MAP:
        remove_caption_and_next_image_after(
            document,
            [caption, APPENDIX_FIGURE_TITLE_MAP[caption]],
            "Appendix A: Non-profit revenue source comparison details",
        )
    for table_caption, caption in [
        ("Table 5-2: Total Revenue by Region", "Figure 5-2: Total Revenue by Region"),
        ("Table 5-3: Program Service Revenue by Region", "Figure 5-3: Program Service Revenue by Region"),
        ("Table 5-4: Total Contributions by Region", "Figure 5-4: Total Contributions by Region"),
        ("Table 5-5: Government Grants Received by Region", "Figure 5-5: Government Grants Received by Region"),
        ("Table 5-6: Federated Campaign Contributions by Region", "Figure 5-6: Federated Campaign Contributions by Region"),
        ("Table 5-7: Related Organization Contributions by Region", "Figure 5-7: Related Organization Contributions by Region"),
        ("Table 5-8: Membership Dues by Region", "Figure 5-8: Membership Dues by Region"),
        ("Table 5-9: Fundraising Event Contributions by Region", "Figure 5-9: Fundraising Event Contributions by Region"),
        ("Table 5-10: Mixed or Unclassified Contributions by Region", "Figure 5-10: Mixed or Unclassified Contributions by Region"),
        ("Table 5-11: Other Revenue by Region", "Figure 5-11: Other Revenue by Region"),
    ]:
        insert_elements_before_caption(
            document,
            [table_caption, APPENDIX_TABLE_TITLE_MAP[table_caption]],
            image_elements_with_new_caption(
                document,
                appendix_figure_elements[caption],
                APPENDIX_FIGURE_TITLE_MAP[caption],
            ),
            after_heading="Appendix A: Non-profit revenue source comparison details",
        )
    retitle_paragraphs_after_heading(
        document,
        {**APPENDIX_FIGURE_TITLE_MAP, **APPENDIX_TABLE_TITLE_MAP},
        after_heading="Appendix A: Non-profit revenue source comparison details",
    )
    replace_first_paragraph_after_heading(
        document,
        "Appendix A: Non-profit revenue source comparison details",
        APPENDIX_INTRO,
    )
    replace_paragraph_starting_with(
        document,
        ["Table 5-13 sums reported dollars", "Table A-1 sums reported dollars"],
        AGGREGATE_TOTALS_NOTE,
    )
    remove_paragraph_starting_with_after(
        document,
        "How to read the comparison tables:",
        after_heading="Appendix A: Non-profit revenue source comparison details",
    )
    insert_elements_before_caption(
        document,
        ["Figure 5-2: Total Revenue by Region", "Figure A-1: Total Revenue by Region"],
        [make_paragraph_element(document, COMPARISON_TABLE_METHOD_NOTE)],
        after_heading="Appendix A: Non-profit revenue source comparison details",
    )
    update_q9_source_table_headers(document)


def is_heading(paragraph: Paragraph, level: int | None = None) -> bool:
    style = paragraph.style.name
    if level is None:
        return style.startswith("Heading")
    return style == f"Heading {level}"


def starts_black_block(text: str) -> str | None:
    if text == IRS_BMF_HEADING:
        return "irs_bmf_data_source"
    if text == "GivingTuesday 990 DataMart":
        return "givingtuesday_data_source"
    if text == NCCS_BMF_HEADING:
        return "nccs_bmf_data_source"
    if text == "Non-profit revenue analysis":
        return "q9_relevant_details"
    if text == "Non-profit revenue source comparison (2022)":
        return "q9_results"
    if text == "Appendix A: Non-profit revenue source comparison details":
        return "q9_appendix"
    return None


def should_end_black_block(block: str | None, paragraph: Paragraph, text: str) -> bool:
    if block is None:
        return False
    if block in {"irs_bmf_data_source", "givingtuesday_data_source", "nccs_bmf_data_source"}:
        return is_heading(paragraph, 3) and text not in {
            IRS_BMF_HEADING,
            "GivingTuesday 990 DataMart",
            NCCS_BMF_HEADING,
        }
    if block == "q9_relevant_details":
        return is_heading(paragraph, 2) and text == "Results of non-profit organization analysis"
    if block == "q9_results":
        return is_heading(paragraph, 3) and text == "Number of Non-profit Organizations from 2022 to 2024"
    if block == "q9_appendix":
        return text.startswith("Table 5-13. Incidence Risk Ratios")
    return False


def color_code_document(doc: Document) -> set[int]:
    """Return the ids of tables that belong to black-added sections."""

    for paragraph in doc.paragraphs:
        set_paragraph_color(paragraph, GRAY)
    for table in doc.tables:
        set_table_color(table, GRAY)

    black_table_ids: set[int] = set()
    active_block: str | None = None

    for item in iter_body_items(doc):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if should_end_black_block(active_block, item, text):
                active_block = None
            start = starts_black_block(text)
            if start is not None:
                active_block = start
            if active_block is not None:
                set_paragraph_color(item, BLACK)
        else:
            if active_block is not None:
                black_table_ids.add(id(item._element))
                set_table_color(item, BLACK)
                shade_header(item)

    color_footnotes_gray(doc)
    return black_table_ids


def color_footnotes_gray(doc: Document) -> None:
    footnotes_part = next(
        (part for part in doc.part.package.parts if str(part.partname) == "/word/footnotes.xml"),
        None,
    )
    if footnotes_part is None:
        return
    root = etree.fromstring(footnotes_part.blob)
    for run in root.findall(".//" + qn("w:r")):
        set_run_color(run, GRAY)
    footnotes_part._blob = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )


def run_color(run) -> str | None:
    if run.font.color.rgb is not None:
        return str(run.font.color.rgb).upper()
    run_properties = run._element.find(qn("w:rPr"))
    color_element = run_properties.find(qn("w:color")) if run_properties is not None else None
    if color_element is not None:
        return (color_element.get(qn("w:val")) or "").upper()
    return None


def paragraph_is_gray(paragraph: Paragraph) -> bool:
    colors = [run_color(run) for run in paragraph.runs if run.text]
    return bool(colors) and all(color == GRAY for color in colors)


def table_is_gray(table: Table) -> bool:
    colors: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                colors.extend(color for run in paragraph.runs if run.text for color in [run_color(run)] if color)
    return bool(colors) and all(color == GRAY for color in colors)


def escape_cell(text: str, *, gray: bool) -> str:
    escaped = html.escape(text.strip()).replace("\n", "<br>")
    escaped = escaped.replace("|", "\\|")
    if gray and escaped:
        return f'<span style="color: #{GRAY};">{escaped}</span>'
    return escaped


def markdown_table(table: Table, *, gray: bool) -> list[str]:
    if not table.rows:
        return []
    rows = [[escape_cell(cell.text, gray=gray) for cell in row.cells] for row in table.rows]
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    lines = [
        "| " + " | ".join(normalized[0]) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def image_extension(partname: str) -> str:
    suffix = Path(partname).suffix.lower()
    return suffix if suffix else ".png"


def save_images_from_paragraph(
    paragraph: Paragraph,
    image_counter: int,
    image_cache: dict[str, str],
) -> tuple[list[str], int]:
    paths: list[str] = []
    for blip in paragraph._element.findall(".//" + qn("a:blip")):
        relationship_id = blip.get(qn("r:embed"))
        if not relationship_id:
            continue
        part = paragraph.part.related_parts[relationship_id]
        cache_key = f"{relationship_id}:{part.partname}"
        if cache_key not in image_cache:
            image_counter += 1
            filename = f"final-report-updated-docx-image-{image_counter:03d}{image_extension(str(part.partname))}"
            path = ASSET_DIR / filename
            path.write_bytes(part.blob)
            image_cache[cache_key] = f"assets/{filename}"
        paths.append(image_cache[cache_key])
    return paths, image_counter


def paragraph_markdown_prefix(paragraph: Paragraph) -> str:
    style = paragraph.style.name
    if style == "Title":
        return "# "
    match = re.fullmatch(r"Heading ([1-6])", style)
    if match:
        return "#" * int(match.group(1)) + " "
    return ""


def paragraph_to_markdown(paragraph: Paragraph) -> str:
    text = paragraph.text.strip()
    if not text:
        return ""
    gray = paragraph_is_gray(paragraph)
    escaped = html.escape(text)
    if gray:
        escaped = f'<span style="color: #{GRAY};">{escaped}</span>'
    return paragraph_markdown_prefix(paragraph) + escaped


def write_markdown(doc: Document) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    lines: list[str] = []
    image_counter = 0
    image_cache: dict[str, str] = {}

    for item in iter_body_items(doc):
        if isinstance(item, Paragraph):
            paragraph_line = paragraph_to_markdown(item)
            if paragraph_line:
                lines.append(paragraph_line)
                lines.append("")
            image_paths, image_counter = save_images_from_paragraph(
                item,
                image_counter,
                image_cache,
            )
            for image_path in image_paths:
                lines.append(f"![drawing]({image_path})")
                lines.append("")
            if not paragraph_line and not image_paths:
                if lines and lines[-1] != "":
                    lines.append("")
        else:
            lines.extend(markdown_table(item, gray=table_is_gray(item)))
            lines.append("")

    while lines and lines[-1] == "":
        lines.pop()
    MD_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    doc = Document(SOURCE_DOCX)
    patch_report_content(doc)
    apply_q9_report_styles(doc)
    format_q9_tables(doc)
    doc.save(SOURCE_DOCX)
    black_tables = color_code_document(doc)
    apply_q9_report_styles(doc)
    format_q9_tables(doc)
    doc.save(COLOR_DOCX)

    color_doc = Document(COLOR_DOCX)
    write_markdown(color_doc)

    with ZipFile(COLOR_DOCX) as zip_file:
        bad = zip_file.testzip()
        if bad is not None:
            raise RuntimeError(f"Corrupt DOCX package member: {bad}")

    print(f"Wrote {COLOR_DOCX}")
    print(f"Wrote {MD_PATH}")
    print(
        f"Paragraphs={len(color_doc.paragraphs)} tables={len(color_doc.tables)} "
        f"images={len(color_doc.inline_shapes)} black_tables={len(black_tables)}"
    )


if __name__ == "__main__":
    main()
