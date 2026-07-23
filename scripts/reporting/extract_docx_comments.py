"""Extract Word comments and their anchored text from a docx file."""

from __future__ import annotations

import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def comment_map(docx_path: Path) -> dict[str, dict]:
    with ZipFile(docx_path) as zf:
        if "word/comments.xml" not in zf.namelist():
            return {}
        root = etree.fromstring(zf.read("word/comments.xml"))

    out: dict[str, dict] = {}
    for comment in root.findall(".//w:comment", NS):
        cid = comment.get(f"{{{W_NS}}}id", "")
        author = comment.get(f"{{{W_NS}}}author", "")
        date = comment.get(f"{{{W_NS}}}date", "")
        texts: list[str] = []
        for node in comment.findall(".//w:t", NS):
            if node.text:
                texts.append(node.text)
            if node.tail:
                texts.append(node.tail)
        out[cid] = {
            "author": author,
            "date": date,
            "text": "".join(texts).strip(),
        }
    return out


def anchored_text(document: Document, comment_id: str) -> str:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        for child in paragraph._element.iter():
            if child.tag == f"{{{W_NS}}}commentRangeStart":
                if child.get(f"{{{W_NS}}}id") == comment_id:
                    parts.append(paragraph.text.strip())
            if child.tag == f"{{{W_NS}}}commentReference":
                if child.get(f"{{{W_NS}}}id") == comment_id and paragraph.text.strip() not in parts:
                    parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for child in paragraph._element.iter():
                        if child.tag == f"{{{W_NS}}}commentRangeStart":
                            if child.get(f"{{{W_NS}}}id") == comment_id:
                                parts.append(paragraph.text.strip())
    return " | ".join(dict.fromkeys(p for p in parts if p))


def extract(docx_path: Path) -> list[dict]:
    comments = comment_map(docx_path)
    if not comments:
        return []
    document = Document(docx_path)
    rows: list[dict] = []
    for cid, meta in sorted(comments.items(), key=lambda item: int(item[0]) if item[0].isdigit() else item[0]):
        rows.append(
            {
                "id": cid,
                "author": meta["author"],
                "date": meta["date"],
                "comment": meta["text"],
                "anchor": anchored_text(document, cid),
            }
        )
    return rows


def main() -> None:
    report_dir = (
        Path(__file__).resolve().parents[2]
        / "docs"
        / "deliverables"
        / "final_report"
    )
    targets = sys.argv[1:] or [
        str(report_dir / "source" / "black_hills_philanthropy_source.docx"),
        str(report_dir / "working" / "black_hills_philanthropy_color_coded.docx"),
    ]
    for name in targets:
        path = report_dir / name if not Path(name).is_absolute() else Path(name)
        print(f"\n=== {path.name} ===")
        if not path.exists():
            print("File not found")
            continue
        rows = extract(path)
        if not rows:
            print("No comments")
            continue
        for index, row in enumerate(rows, 1):
            date = row["date"][:10] if row["date"] else ""
            print(f"[{index}] id={row['id']} author={row['author']} date={date}")
            if row["anchor"]:
                print(f"Anchor: {row['anchor'][:240]}")
            print(f"Comment: {row['comment']}")
            print("-" * 72)
        print(f"Total comments: {len(rows)}")


if __name__ == "__main__":
    main()
