"""Quick audit of the color-coded final report docx."""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path
from zipfile import ZipFile

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree
from sync_final_report_black_sections import (
    ADDED_TABLE_HEADER_FILL,
    APPENDIX_START,
    BODY_FONT_SIZE_PT,
    HEADING_4_TITLES,
    HEADER_FONT_SIZE_PT,
    LEGACY_SOURCE_TABLE_HEADERS,
    OVERVIEW_TABLE_HEADERS,
    PRIOR_SOURCE_TABLE_HEADERS,
    REVENUE_DEFINITION_TABLE_HEADERS,
    REVENUE_DEFINITION_TABLE_ROWS,
    REPORT_VARIABLES,
    REPORT_REGION_ORDER,
    RESULTS_START,
    SOURCE_TABLE_HEADERS,
    SIGNIFICANT_TABLE_HEADERS,
    load_overview_rows,
    load_aggregate_rows,
    load_source_rows,
    paragraph_has_image,
    paragraph_is_gray,
    report_values,
)

DOCX = Path(
    os.environ.get(
        "FINAL_REPORT_DOCX",
        Path(__file__).resolve().parents[1]
        / "docs/final_report/Philanthropic characteristics of the Black Hills Area and co color-coded.docx",
    )
)
BASELINE_DOCX = Path(
    os.environ.get(
        "FINAL_REPORT_BASELINE_DOCX",
        Path(os.environ.get("TEMP", ".")) / "final_report_color_baseline.docx",
    )
)
SOURCE_DOCX = Path(__file__).resolve().parents[1] / (
    "docs/final_report/Philanthropic characteristics of the Black Hills Area and co (UPDATED).docx"
)
EXPECTED_ASSET_HASHES_1_TO_20 = {
    1: "F53698FD5CF0B1F7CB0DFC22C53D12A2F16FBD2BE566B23ECA5CE79C77773987",
    2: "E6B2897E08404029B8531810905C7AD495DAD14FB3D33ADD9D2A034A02D9900F",
    3: "F922CFFD626EB783557378D3C43C2FB1EB97D763CC61F973E73D3EB5B04C11AC",
    4: "3F48ADF656BF7EC71D9C086956E532D55D3C866688156C99B76B97022282B2E0",
    5: "8C0311DECADB8BB70FE8B1C93162243F21B222EA69857D63E68939E2C2A702D8",
    6: "736A8958BACA3110A122ED8771C1C17BD87A8FEB8DF63EF6434DD77CABAF75E2",
    7: "D45100DC3D5252AE2CC02FC75EB0BBAC5E1CD3AA28739A6999FBE9A4658B4017",
    8: "13636E9B9118F8F256572CDFE55E9672EBAF4BABA242F7C4A4E1249233DDF993",
    9: "9E0C74910239BCC78C472233E50C21181D111E2F562B6678234D69A375D07404",
    10: "9A4CEF15CC155C314151EDE504548B27E74E218BA09A1719BC1F482A37D3403E",
    11: "9A09F4493EFA5F8E89998E9EE2506338179E203625F7BAAADFF905A7A91A994D",
    12: "B392D8151A7D6BA69DAD573192997B4BA1DCEF62C2E2E920571FD09290254BBD",
    13: "B1045157AE67EF9A84E6C047DF1E4E1932658F8598A7A4680AFD62DA76CBEE2B",
    14: "AEA3E204AA0DDA1E3A6343EA31D2C05F9E72FC8571270C0C011E7E653445DAE5",
    15: "40E101F686D2B1E0C8D8EBB032A16EF6DCCA6C325D25918D18C8BACCB12FC3E0",
    16: "D2C2C21446703256901ABB23A1F606612D6CCF95C5FCAEA2507F1FB92719E8D8",
    17: "5A6B3FCA24ED19D11264671C443EB27D125D70BC6A9F8D31B12E3AF42C904A19",
    18: "5835241CA1147D634F30C460732B3552AD577FF307655E43CC6A23EF46559CAE",
    19: "7DB94619961823340A7E130701E69D12CF98891C4C395BE03BFAC12633532066",
    20: "8729F81A52030486EBB19701B34284EB504E46A84896858D7337F0DF5669E2FD",
}


def gray_paragraphs(document: Document):
    return [paragraph for paragraph in document.paragraphs if paragraph_is_gray(paragraph)]


def is_arial_run(run, *, size_pt: float) -> bool:
    name = run.font.name
    size = run.font.size.pt if run.font.size is not None else None
    return name == "Arial" and size == size_pt


def expected_black_paragraph_size(paragraph: Paragraph) -> float:
    text = paragraph.text.strip()
    if paragraph.style.name.startswith("Heading") or text in HEADING_4_TITLES:
        return float(HEADER_FONT_SIZE_PT)
    return float(BODY_FONT_SIZE_PT)


def body_items(document: Document):
    for element in document.element.body.iterchildren():
        if element.tag.endswith("}p"):
            yield Paragraph(element, document._body)
        elif element.tag.endswith("}tbl"):
            yield Table(element, document._body)


def main() -> None:
    doc = Document(DOCX)
    full = "\n".join(p.text for p in doc.paragraphs)
    issues: list[str] = []

    appendix_headings = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith(APPENDIX_START)
    ]
    if len(appendix_headings) != 1:
        issues.append(f"Appendix headings: {len(appendix_headings)} (expect 1)")
    if full.count("Participation and dollar amounts") > 1:
        issues.append("Duplicate participation paragraph")
    datamart_headings = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip() == "GivingTuesday 990 DataMart"
    ]
    if len(datamart_headings) != 1:
        issues.append(f"GivingTuesday DataMart headings: {len(datamart_headings)} (expect 1)")

    for stale in (
        "Black Hills share of aggregate",
        "Typical amount, Black Hills",
        "in plain terms",
    ):
        if stale in full:
            issues.append(f"Stale text: {stale!r}")

    sig = [
        table
        for table in doc.tables
        if [cell.text.strip() for cell in table.rows[0].cells] == SIGNIFICANT_TABLE_HEADERS
    ]
    overview = [
        table
        for table in doc.tables
        if [cell.text.strip() for cell in table.rows[0].cells] == OVERVIEW_TABLE_HEADERS
    ]
    app = [
        table
        for table in doc.tables
        if [cell.text.strip() for cell in table.rows[0].cells] == SOURCE_TABLE_HEADERS
    ]
    legacy_app = [
        table
        for table in doc.tables
        if [cell.text.strip() for cell in table.rows[0].cells] == LEGACY_SOURCE_TABLE_HEADERS
    ]
    prior_app = [
        table
        for table in doc.tables
        if [cell.text.strip() for cell in table.rows[0].cells] == PRIOR_SOURCE_TABLE_HEADERS
    ]
    definitions = [
        table
        for table in doc.tables
        if [cell.text.strip() for cell in table.rows[0].cells] == REVENUE_DEFINITION_TABLE_HEADERS
    ]
    if sig:
        issues.append(f"Significant-results tables remain: {len(sig)}")
    if legacy_app:
        issues.append(f"Appendix tables with Result column remain: {len(legacy_app)}")
    if prior_app:
        issues.append(f"Appendix tables without reporting rates remain: {len(prior_app)}")
    if len(definitions) != 1:
        issues.append(f"Revenue definition tables: {len(definitions)} (expect 1)")
    elif [[cell.text.strip() for cell in row.cells] for row in definitions[0].rows[1:]] != REVENUE_DEFINITION_TABLE_ROWS:
        issues.append("Revenue definition table values do not match the expected definitions")
    if len(overview) != 2:
        issues.append(f"Overview/aggregate value tables: {len(overview)} (expect two 6-column tables)")
    if len(app) != 10:
        issues.append(f"Appendix tables: {len(app)} (expect 10)")
    summary, values = report_values()
    if len(overview) == 2:
        median_table, aggregate_table_doc = overview
        if len(median_table.rows) != 11:
            issues.append(f"Median overview table rows: {len(median_table.rows)} (expect 11)")
        if len(aggregate_table_doc.rows) != 11:
            issues.append(f"Aggregate overview table rows: {len(aggregate_table_doc.rows)} (expect 11)")
        actual_median = [[cell.text.strip() for cell in row.cells] for row in median_table.rows[1:]]
        if actual_median != load_overview_rows(values):
            issues.append("Overview table values do not match the presentation summary CSV")
        actual_aggregate = [[cell.text.strip() for cell in row.cells] for row in aggregate_table_doc.rows[1:]]
        if actual_aggregate != load_aggregate_rows():
            issues.append("Aggregate table values do not match client_raw_level_region_summary.csv")
    for index, table in enumerate(app, start=2):
        regions = [row.cells[0].text.strip() for row in table.rows[1:]]
        if regions != REPORT_REGION_ORDER:
            issues.append(f"Table 5-{index} region order: {regions}")
    if len(app) == 10:
        for table, variable in zip(app, REPORT_VARIABLES):
            actual = [[cell.text.strip() for cell in row.cells] for row in table.rows[1:]]
            if actual != load_source_rows(summary, values, variable):
                issues.append(f"Source table values do not match CSV: {variable}")
            if any("Reference" in cell.text for row in table.rows for cell in row.cells):
                issues.append(f"Reference placeholder remains in source table: {variable}")
            if any(
                "Not applicable (baseline)" in cell.text
                for row in table.rows
                for cell in row.cells
            ):
                issues.append(f"Long baseline placeholder remains in source table: {variable}")
            if table.rows[1].cells[4].text.strip() != "NA" or table.rows[1].cells[5].text.strip() != "NA":
                issues.append(f"Black Hills comparison cells are not NA: {variable}")

    empty = sum(
        1
        for p in doc.paragraphs
        if not p.text.strip() and not paragraph_is_gray(p) and not paragraph_has_image(p)
    )
    if empty:
        issues.append(f"Empty black paragraphs: {empty}")

    if full.count("Figure 5-1:") != 0:
        issues.append(f"Figure 5-1 title count: {full.count('Figure 5-1:')} (expect 0)")
    for n in range(2, 12):
        if full.count(f"Figure 5-{n}:") != 1:
            issues.append(f"Figure 5-{n} title count: {full.count(f'Figure 5-{n}:')}")
    for n in range(1, 14):
        if full.count(f"Table 5-{n}:") != 1:
            issues.append(f"Table 5-{n} title count: {full.count(f'Table 5-{n}:')}")

    paragraphs = doc.paragraphs
    results_i = next(
        (i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip().startswith(RESULTS_START)),
        None,
    )
    conclusions_i = next(
        (
            i
            for i, paragraph in enumerate(paragraphs)
            if paragraph_is_gray(paragraph) and paragraph.text.strip() == "Conclusions"
        ),
        None,
    )
    appendix_i = next(
        (i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip().startswith(APPENDIX_START)),
        None,
    )
    if results_i is None or conclusions_i is None:
        issues.append("Could not locate the added results and gray Conclusions boundaries")
    else:
        main_paragraphs = paragraphs[results_i:conclusions_i]
        main_text = "\n".join(paragraph.text for paragraph in main_paragraphs).lower()
        if "appendix table 5-1 provides the exact median reported dollars" not in main_text:
            issues.append("Missing direct reference to Appendix Table 5-1 in the main body")
        if "appendix table 5-13 provides the corresponding regional totals" not in main_text:
            issues.append("Missing direct reference to Appendix Table 5-13 in the main body")
        if "figure 5-1" in main_text or "overview figure" in main_text:
            issues.append("Figure 5-1 language remains in the main body")
        prohibited = (
            "p-value",
            "p =",
            "p <",
            "p &lt;",
            "significant",
            "not significant",
            "statistically meaningful",
            "permutation test",
            "confidence interval",
            "95% ci",
        )
        for phrase in prohibited:
            if phrase in main_text:
                issues.append(f"Main-body inferential language remains: {phrase!r}")
        main_images = sum(paragraph_has_image(paragraph) for paragraph in main_paragraphs)
        if main_images != 0:
            issues.append(f"Added main-body figures: {main_images} (expect 0)")

        results_element = paragraphs[results_i]._element
        conclusions_element = paragraphs[conclusions_i]._element
        in_main = False
        main_tables = 0
        for item in body_items(doc):
            if isinstance(item, Paragraph) and item._element is results_element:
                in_main = True
            if isinstance(item, Paragraph) and item._element is conclusions_element:
                in_main = False
                break
            if in_main and isinstance(item, Table):
                main_tables += 1
        if main_tables:
            issues.append(f"Added main-body tables: {main_tables} (expect 0)")

    if conclusions_i is not None and appendix_i is not None and appendix_i <= conclusions_i:
        issues.append("Appendix heading is not after the gray Conclusions section")

    if "We ran the analysis both with and without 25 records" not in full:
        issues.append("Missing with-and-without special-organization sensitivity explanation")
    if "did not reverse the direction of any of the 40 Black Hills-to-benchmark comparisons" not in full:
        issues.append("Missing sensitivity result for the 40 comparisons")
    if "local organizations may be accessing some funding channels but receiving smaller amounts" not in full:
        issues.append("Missing explicit connection from results to the Foundation interpretation")

    if appendix_i is not None:
        table_5_1_i = next(
            (
                i
                for i, paragraph in enumerate(paragraphs)
                if paragraph.text.strip().startswith("Table 5-1:")
            ),
            None,
        )
        figure_5_2_i = next(
            (
                i
                for i, paragraph in enumerate(paragraphs)
                if paragraph.text.strip().startswith("Figure 5-2:")
            ),
            None,
        )
        if table_5_1_i is None or table_5_1_i <= appendix_i:
            issues.append("Table 5-1 is not in the appendix")
        table_5_13_i = next(
            (
                i
                for i, paragraph in enumerate(paragraphs)
                if paragraph.text.strip().startswith("Table 5-13:")
            ),
            None,
        )
        if table_5_13_i is None or table_5_13_i <= appendix_i:
            issues.append("Table 5-13 is not in the appendix")
        if figure_5_2_i is not None and table_5_1_i is not None and table_5_1_i >= figure_5_2_i:
            issues.append("Table 5-1 does not precede Figure 5-2 in the appendix")
        if (
            figure_5_2_i is not None
            and table_5_13_i is not None
            and table_5_13_i >= figure_5_2_i
        ):
            issues.append("Table 5-13 does not precede Figure 5-2 in the appendix")
        table_5_12_i = next(
            (
                i
                for i, paragraph in enumerate(paragraphs)
                if paragraph.text.strip().startswith("Table 5-12:")
            ),
            None,
        )
        if table_5_12_i is None or table_5_12_i <= appendix_i:
            issues.append("Table 5-12 is not in the appendix")

    imgs_total = len(doc.inline_shapes)
    if imgs_total != 30:
        issues.append(f"Total embedded charts: {imgs_total} (expect 30)")

    for paragraph in doc.paragraphs:
        if paragraph_is_gray(paragraph) or paragraph_has_image(paragraph) or not paragraph.text.strip():
            continue
        expected_size = expected_black_paragraph_size(paragraph)
        for run in paragraph.runs:
            if run.text and not is_arial_run(run, size_pt=expected_size):
                issues.append(f"Unexpected black run font: {paragraph.text[:80]!r}")
                break
        if re.match(r"^(Figure|Table) 5-\d+:", paragraph.text.strip()):
            if not all(run.bold for run in paragraph.runs if run.text):
                issues.append(f"Unbolded title: {paragraph.text!r}")
            if any(run.underline for run in paragraph.runs):
                issues.append(f"Underlined title: {paragraph.text!r}")
            p_pr = paragraph._element.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:pBdr")) is not None:
                issues.append(f"Bordered title: {paragraph.text!r}")

    for table in [*overview, *app, *definitions]:
        header_fills = []
        for cell in table.rows[0].cells:
            shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            header_fills.append(shading.get(qn("w:fill")) if shading is not None else None)
        if any(fill != ADDED_TABLE_HEADER_FILL for fill in header_fills):
            issues.append(f"Added table header is not light blue: {header_fills}")
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text and not is_arial_run(run, size_pt=BODY_FONT_SIZE_PT):
                            issues.append(f"Unexpected table run font: {run.text!r}")
                            break

    if BASELINE_DOCX.exists():
        baseline = Document(BASELINE_DOCX)
        before = [paragraph._element.xml for paragraph in gray_paragraphs(baseline)]
        after = [paragraph._element.xml for paragraph in gray_paragraphs(doc)]
        if before != after:
            issues.append("Gray paragraph XML changed from the pre-sync baseline")
        with ZipFile(BASELINE_DOCX) as baseline_zip, ZipFile(DOCX) as current_zip:
            baseline_media = {
                name for name in baseline_zip.namelist() if name.startswith("media/")
            }
            current_media = {
                name for name in current_zip.namelist() if name.startswith("media/")
            }
            shared = baseline_media & current_media
            changed = [
                name for name in shared
                if baseline_zip.read(name) != current_zip.read(name)
            ]
            added = current_media - baseline_media
            removed = baseline_media - current_media
            if len(current_media) != 30 or changed or added or len(removed) > 1:
                issues.append(
                    f"Embedded media audit: {len(current_media)} current, {len(changed)} changed, "
                    f"{len(added)} added, and {len(removed)} removed "
                    "(expect 30 current, 0 changed, 0 added, and at most 1 removed)"
                )
    elif SOURCE_DOCX.exists():
        source_text = [paragraph.text for paragraph in Document(SOURCE_DOCX).paragraphs]
        gray_text = [paragraph.text for paragraph in gray_paragraphs(doc)]
        position = 0
        for text in source_text:
            if position < len(gray_text) and text == gray_text[position]:
                position += 1
        if position != len(gray_text):
            issues.append("Gray wording is not an unchanged subsequence of the teammate source DOCX")

    with ZipFile(DOCX) as current_zip:
        root = etree.fromstring(current_zip.read("word/footnotes.xml"))
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        target_prefixes = (
            "State tax provisions reflect rules in effect for the 2022 tax year",
            "Both r and rs are correlation coefficients",
        )
        for prefix in target_prefixes:
            matches = []
            for footnote in root.xpath("//w:footnote", namespaces=namespace):
                text = "".join(
                    footnote.xpath(".//w:t/text()", namespaces=namespace)
                ).strip()
                if text.startswith(prefix):
                    matches.append(footnote)
            if len(matches) != 1:
                issues.append(f"Teammate footnote count for {prefix!r}: {len(matches)}")
                continue
            for run in matches[0].xpath(".//w:r", namespaces=namespace):
                text = "".join(run.xpath(".//w:t/text()", namespaces=namespace))
                if not text:
                    continue
                colors = run.xpath("./w:rPr/w:color/@w:val", namespaces=namespace)
                if colors != ["666666"]:
                    issues.append(f"Teammate footnote is not gray: {prefix!r}")
                    break

    import hashlib

    asset_dir = DOCX.parent / "updated_report_assets"
    for index, expected_hash in EXPECTED_ASSET_HASHES_1_TO_20.items():
        path = asset_dir / f"updated-report-image-{index:02d}.png"
        actual_hash = hashlib.sha256(path.read_bytes()).hexdigest().upper()
        if actual_hash != expected_hash:
            issues.append(f"Protected image {index:02d} changed")

    print(f"Auditing: {DOCX.name}")
    print(f"Issues: {len(issues)}")
    for item in issues:
        print(f"  - {item}")
    if not issues:
        print("  - No automated issues found")


if __name__ == "__main__":
    main()
