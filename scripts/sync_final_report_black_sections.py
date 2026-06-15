"""Sync black (non-gray) paragraphs from final report markdown into color-coded docx files."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

REVENUE_DEFINITION_COLUMN_FRACTIONS = (0.17, 0.33, 0.50)
ADDED_TABLE_HEADER_FILL = "DEEAF0"
BODY_FONT_SIZE_PT = 8
HEADER_FONT_SIZE_PT = 13

REPO_ROOT = Path(__file__).resolve().parents[1]
MD_PATH = REPO_ROOT / "docs/final_report/Philanthropic characteristics of the Black Hills Area and co (1).md"
DOCX_PATH = REPO_ROOT / "docs/final_report/Philanthropic characteristics of the Black Hills Area and co color-coded.docx"
PAIRWISE_CSV = (
    REPO_ROOT
    / "python/analysis/revenue_sources_black_hills/results/tables/client_2022_pairwise_presentation_summary.csv"
)
RAW_LEVEL_CSV = (
    REPO_ROOT
    / "python/analysis/revenue_sources_black_hills/results/tables/client_raw_level_region_summary.csv"
)
SIGNIFICANT_TABLE_HEADERS = [
    "Revenue source",
    "Benchmark region",
    "Direction",
    "Black Hills median",
    "Benchmark median",
    "Median gap (95% CI)",
    "p-value",
]
APPENDIX_TABLE_HEADERS = SIGNIFICANT_TABLE_HEADERS[1:]
OVERVIEW_TABLE_HEADERS = [
    "Revenue source",
    "Black Hills",
    "Billings",
    "Flagstaff",
    "Missoula",
    "Sioux Falls",
]
SOURCE_TABLE_HEADERS = [
    "Region",
    "Reporting rate",
    "Positive reporters (n)",
    "Median reported dollars",
    "Black Hills minus benchmark region (95% CI)",
    "Pairwise p-value vs Black Hills",
]
PRIOR_SOURCE_TABLE_HEADERS = [
    "Region",
    "Positive reporters (n)",
    "Median reported dollars",
    "Black Hills minus region (95% CI)",
    "p-value",
]
LEGACY_SOURCE_TABLE_HEADERS = [
    *PRIOR_SOURCE_TABLE_HEADERS,
    "Result",
]
REPORT_REGION_ORDER = ["Black Hills", "Billings", "Flagstaff", "Missoula", "Sioux Falls"]
REPORT_VARIABLES = [
    "total_revenue",
    "program_service_revenue",
    "total_contributions",
    "government_grants_received",
    "federated_campaigns",
    "related_org_contributions",
    "membership_dues",
    "fundraising_events_contributions",
    "mixed_unclassified_contributions",
    "residual_other_revenue",
]
REPORT_DISPLAY_LABELS = {
    "total_revenue": "Total revenue",
    "program_service_revenue": "Program service revenue",
    "total_contributions": "Total contributions",
    "government_grants_received": "Government grants received",
    "federated_campaigns": "Federated campaign contributions",
    "related_org_contributions": "Related organization contributions",
    "membership_dues": "Membership dues",
    "fundraising_events_contributions": "Fundraising event contributions",
    "mixed_unclassified_contributions": "Mixed or unclassified contributions",
    "residual_other_revenue": "Other revenue",
}
REPORT_TITLE_LABELS = {
    "total_revenue": "Total Revenue",
    "program_service_revenue": "Program Service Revenue",
    "total_contributions": "Total Contributions",
    "government_grants_received": "Government Grants Received",
    "federated_campaigns": "Federated Campaign Contributions",
    "related_org_contributions": "Related Organization Contributions",
    "membership_dues": "Membership Dues",
    "fundraising_events_contributions": "Fundraising Event Contributions",
    "mixed_unclassified_contributions": "Mixed or Unclassified Contributions",
    "residual_other_revenue": "Other Revenue",
}
FIGURE_TITLES = {
    **{
        variable: f"Figure 5-{index}: {REPORT_TITLE_LABELS[variable]} by Region"
        for index, variable in enumerate(REPORT_VARIABLES, start=2)
    },
}
TABLE_TITLES = {
    "overview": "Table 5-1: Median Reported Dollars by Revenue Source and Region",
    "aggregate": "Table 5-13: Total Reported Dollars by Revenue Source and Region (Aggregate)",
    "definitions": "Table 5-12: Revenue Source Definitions",
    **{
        variable: f"Table 5-{index}: {REPORT_TITLE_LABELS[variable]} by Region"
        for index, variable in enumerate(REPORT_VARIABLES, start=2)
    },
}
REPORT_IMAGE_PATHS = [
    REPO_ROOT / "docs" / "final_report" / "updated_report_assets" / f"updated-report-image-{index:02d}.png"
    for index in range(22, 32)
]
APPENDIX_INTRO = (
    "The appendix provides the complete 2022 revenue-source comparison details. Table 5-1 lists the median "
    "reported dollars among organizations with a positive amount for each source. Table 5-13 lists the "
    "corresponding regional totals summed across all organizations in the analysis file. Figures 5-2 through "
    "5-11 show each source separately, and their paired tables provide reporting rates, reporter counts, "
    "medians, median gaps, confidence intervals, and p-values. The comparison columns apply only to each "
    "benchmark region versus Black Hills; no Black Hills-only p-value or median-gap confidence interval "
    "exists, so those Black Hills cells are marked NA. "
    "Table 5-12 defines the revenue sources. Rare revenue sources have fewer reporting organizations; see "
    "the sample-size discussion in the non-profit revenue analysis section."
)
AGGREGATE_TABLE_NOTE = (
    "Table 5-13 sums reported dollars for each revenue source and region over that source's eligible "
    "reporting universe in the 2022 analysis file. Totals for total revenue, total contributions, mixed or "
    "unclassified contributions, and other revenue include all organizations in the file. Government grants, "
    "federated campaigns, related-organization contributions, membership dues, and fundraising event "
    "contributions include Form 990 filers only; 990-EZ and 990-PF filers do not report those Part VIII "
    "sub-lines and are excluded from those rows, not counted as zero. Within Form 990, a blank sub-line is "
    "treated as zero. Program service revenue includes organizations for which that line is reported on the "
    "filed form. Regional totals reflect both universe size and report size; a few very large filers can "
    "dominate the sum. Table 5-1 reports medians among positive reporters only and is the basis for the "
    "pairwise comparisons in Tables 5-2 through 5-11."
)
APPENDIX_FINAL_NOTE = (
    "The full client presentation and supporting analysis files include reporting rates and additional "
    "source-by-source interpretation."
)
FINDINGS_INTRO_PREFIX = (
    "Among organizations that reported a positive amount for the source, Black Hills showed"
)
FINDINGS_AFTER_PREFIX = "For total revenue and total contributions"
STALE_PRE_PREFIXES = (
    "The GivingTuesday 990 DataMart is based on public IRS Form 990-family filings",
    "The current project file covers tax year 2022 for the Black Hills and benchmark-region counties.",
    "The GivingTuesday data contain financial detail that is not available",
    "Hospitals, universities, and political organizations are flagged in the project data",
    "Because these are public filing records, a small number of very large organizations can affect regional totals",
    "The revenue categories mean the following",
    "Program service revenue is money an organization earns",
)
STALE_RESULTS_PREFIXES = (
    "These data do not show why the revenue mix differs",
    "These findings describe reported revenue patterns by region in 2022. They do not establish why",
)

GRAY = "666666"
RESULTS_START = "We compared whether non-profit revenue sources differ"
REVENUE_COMPARISON_HEADING = "Non-profit revenue source comparison (2022)"
APPENDIX_START = "Appendix: Non-profit revenue source comparison details"
DATA_SOURCES_HEADING = "GivingTuesday 990 DataMart"
RELEVANT_DETAILS_HEADING = "Non-profit revenue analysis"
HEADING_4_TITLES = {
    "How many organizations could we compare?",
    "What the 2022 revenue-source comparison shows",
    REVENUE_COMPARISON_HEADING,
    DATA_SOURCES_HEADING,
    RELEVANT_DETAILS_HEADING,
}
REVENUE_DEFINITIONS_INTRO_PREFIX = "The revenue sources analyzed in this section are defined below"
REVENUE_DEFINITION_TABLE_HEADERS = [
    "Revenue source",
    "What it measures",
    "How it is reported on IRS filings",
]
REVENUE_DEFINITION_TABLE_ROWS = [
    [
        "Total revenue",
        "All reported revenue for the organization in tax year 2022.",
        "Form 990, Form 990-EZ, and Form 990-PF totals.",
    ],
    [
        "Total contributions",
        "Total reported contributions, gifts, grants, and similar amounts.",
        "Form 990, Form 990-EZ, and Form 990-PF totals.",
    ],
    [
        "Program service revenue",
        "Money earned from services, programs, fees, or sales related to the organization's mission.",
        "Form 990 and Form 990-EZ; not comparable on Form 990-PF.",
    ],
    [
        "Government grants received",
        "Funds reported from government sources.",
        "Form 990 Part VIII Line 1e; comparisons use Form 990 filers only.",
    ],
    [
        "Federated campaign contributions",
        "Support reported through federated campaigns (for example, United Way-style allocations).",
        "Form 990 Part VIII Line 1a; comparisons use Form 990 filers only.",
    ],
    [
        "Related organization contributions",
        "Contributions from related organizations or affiliates.",
        "Form 990 Part VIII Line 1d; comparisons use Form 990 filers only.",
    ],
    [
        "Membership dues",
        "Membership dues reported on the return.",
        "Form 990 Part VIII Line 1b; comparisons use Form 990 filers only.",
    ],
    [
        "Fundraising event contributions",
        "Contributions tied to fundraising events.",
        "Form 990 Part VIII Line 1c; comparisons use Form 990 filers only.",
    ],
    [
        "Mixed / unclassified contributions",
        "Contribution dollars the forms do not break down cleanly by donor type.",
        "Form 990 Line 1f plus 990-EZ/PF contribution totals that cannot be split into Part VIII lines.",
    ],
    [
        "Other revenue",
        "Remaining revenue after program service revenue and the contribution components above.",
        "Calculated as total revenue minus those components.",
    ],
]
REVENUE_DEFINITIONS_CAVEAT_PREFIX = (
    "The data cannot cleanly separate individual giving from foundation grants"
)
APPENDIX_H3_TITLES: set[str] = set()


def is_gray_md_line(line: str) -> bool:
    stripped = line.strip()
    return bool(stripped) and "#666666" in stripped


def extract_black_paragraphs_from_md(md_text: str) -> list[str]:
    paragraphs: list[str] = []
    in_table = False
    for raw_line in md_text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            in_table = False
            continue
        if stripped.startswith("|"):
            in_table = True
            continue
        if in_table or stripped.startswith("!["):
            continue
        if is_gray_md_line(stripped):
            continue
        if stripped.startswith("*") and stripped.endswith("*") and stripped.count("*") >= 2:
            paragraphs.append(stripped.strip("*").strip())
            continue
        if stripped.startswith("#### "):
            paragraphs.append(stripped.removeprefix("#### ").strip())
            continue
        if stripped.startswith("### "):
            paragraphs.append(stripped.removeprefix("### ").strip())
            continue
        if stripped.startswith("## ") and not is_gray_md_line(stripped):
            paragraphs.append(stripped.removeprefix("## ").strip())
            continue
        if stripped.startswith(("**Figure ", "**Table ")) and stripped.endswith("**"):
            paragraphs.append(stripped.replace("**", "").strip())
            continue
        if stripped.startswith("*Figure "):
            paragraphs.append(stripped.strip("*").strip())
            continue
        if stripped.startswith("- "):
            paragraphs.append(stripped)
            continue
        paragraphs.append(stripped)
    return paragraphs


def paragraph_is_gray(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        if run.font.color.rgb is not None and str(run.font.color.rgb).lower().replace("#", "") == GRAY:
            return True
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is not None:
            color_el = r_pr.find(qn("w:color"))
            if color_el is not None:
                if (color_el.get(qn("w:val")) or "").lower() == GRAY:
                    return True
    return False


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def set_paragraph_style(paragraph: Paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        paragraph.style = "Normal"


def style_for_black_text(text: str, *, in_appendix: bool = False) -> str:
    stripped = text.strip()
    if stripped.startswith(APPENDIX_START):
        return "Heading 2"
    if stripped in HEADING_4_TITLES:
        return "Heading 3"
    if in_appendix and stripped in APPENDIX_H3_TITLES:
        return "Heading 3"
    return "Normal"


def set_paragraph_text_black(
    paragraph: Paragraph,
    text: str,
    *,
    style: str | None = None,
) -> None:
    for child in list(paragraph._element):
        if child.tag.endswith("}r"):
            paragraph._element.remove(child)
    if style is not None:
        set_paragraph_style(paragraph, style)
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str,
    *,
    style: str | None = None,
) -> Paragraph:
    resolved_style = style if style is not None else style_for_black_text(text)
    new_p = deepcopy(paragraph._element)
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)
    p_pr = new_p.find(qn("w:pPr"))
    if p_pr is not None:
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            p_pr.remove(p_style)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_paragraph_text_black(new_para, text, style=resolved_style)
    return new_para


def insert_paragraph_before(
    paragraph: Paragraph,
    text: str,
    *,
    style: str | None = None,
) -> Paragraph:
    resolved_style = style if style is not None else style_for_black_text(text)
    new_p = deepcopy(paragraph._element)
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    p_pr = new_p.find(qn("w:pPr"))
    if p_pr is not None:
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            p_pr.remove(p_style)
    paragraph._element.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_paragraph_text_black(new_para, text, style=resolved_style)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def remove_image_paragraph(paragraph: Paragraph) -> None:
    relationship_ids = {
        blip.get(qn("r:embed"))
        for blip in paragraph._element.findall(".//" + qn("a:blip"))
        if blip.get(qn("r:embed"))
    }
    remove_paragraph(paragraph)
    for relationship_id in relationship_ids:
        paragraph.part.drop_rel(relationship_id)


def remove_aggregate_share_table(doc: Document) -> None:
    for table in list(doc.tables):
        if len(table.rows) < 2:
            continue
        header = " ".join(cell.text for cell in table.rows[0].cells).lower()
        if "black hills share of aggregate" in header:
            table._element.getparent().remove(table._element)


def content_width_emu(doc: Document) -> int:
    section = doc.sections[0]
    return int(section.page_width - section.left_margin - section.right_margin)


def emu_to_twips(emu: int) -> int:
    return int(round(emu * 1440 / 914400))


def set_cell_width_twips(cell, width_twips: int) -> None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))


def set_table_full_page_width(
    table: Table,
    doc: Document,
    column_fractions: tuple[float, ...],
) -> None:
    """Span the printable area and use fixed column widths (template cells are often narrow)."""
    total_emu = content_width_emu(doc)
    total_twips = emu_to_twips(total_emu)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for tag in ("w:tblW", "w:tblLayout", "w:tblInd"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_pr.append(tbl_w)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    if len(column_fractions) != len(table.columns):
        raise ValueError("column_fractions must match table column count")

    assigned_twips = 0
    col_twips: list[int] = []
    for idx, fraction in enumerate(column_fractions):
        if idx == len(column_fractions) - 1:
            col_twips.append(total_twips - assigned_twips)
        else:
            width = int(round(total_twips * fraction))
            col_twips.append(width)
            assigned_twips += width

    for col_idx, width_twips in enumerate(col_twips):
        width_emu = int(round(width_twips * 914400 / 1440))
        table.columns[col_idx].width = width_emu
        for row in table.rows:
            set_cell_width_twips(row.cells[col_idx], width_twips)


def remove_revenue_definition_tables(doc: Document) -> None:
    for table in list(doc.tables):
        if len(table.rows) < 2:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header == REVENUE_DEFINITION_TABLE_HEADERS:
            table._element.getparent().remove(table._element)


def insert_definition_table_after(
    doc: Document,
    paragraph: Paragraph,
    headers: list[str],
    rows: list[list[str]],
) -> Table:
    template = find_appendix_template_table(doc)
    template_header = template.rows[0].cells
    template_data = template.rows[1].cells
    style_map = [0, 1, 1]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = template.style

    all_rows = [headers, *rows]
    for row_idx, row_values in enumerate(all_rows):
        template_row = template_header if row_idx == 0 else template_data
        for col_idx, value in enumerate(row_values):
            set_cell_text_like_template(
                table.rows[row_idx].cells[col_idx],
                value,
                template_row[style_map[col_idx]],
            )

    tbl_element = table._tbl
    doc.element.body.remove(tbl_element)
    paragraph._element.addnext(tbl_element)
    set_table_full_page_width(table, doc, REVENUE_DEFINITION_COLUMN_FRACTIONS)
    return table


def upsert_revenue_definitions_table(doc: Document) -> None:
    intro_para: Paragraph | None = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(REVENUE_DEFINITIONS_INTRO_PREFIX):
            intro_para = paragraph
            break
    if intro_para is None:
        return

    remove_following_revenue_definition_table(intro_para)

    to_remove: list[Paragraph] = []
    found_intro = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(REVENUE_DEFINITIONS_INTRO_PREFIX):
            found_intro = True
            continue
        if not found_intro:
            continue
        if text.startswith(REVENUE_DEFINITIONS_CAVEAT_PREFIX):
            break
        if next_element_is_table(paragraph):
            break
        if text.startswith("A small number of very large organizations"):
            break

    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)

    insert_definition_table_after(
        doc,
        intro_para,
        REVENUE_DEFINITION_TABLE_HEADERS,
        REVENUE_DEFINITION_TABLE_ROWS,
    )


def remove_following_revenue_definition_table(paragraph: Paragraph) -> None:
    next_element = paragraph._element.getnext()
    if next_element is not None and next_element.tag.endswith("tbl"):
        paragraph._element.getparent().remove(next_element)


def remove_significant_results_tables(doc: Document) -> None:
    for table in list(doc.tables):
        if len(table.rows) != 11:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header and header[0] == "Revenue source" and (
            "Black Hills median" in header or "Typical amount" in " ".join(header)
        ):
            table._element.getparent().remove(table._element)


def paragraph_following_element(paragraph: Paragraph) -> Paragraph | None:
    next_element = paragraph._element.getnext()
    while next_element is not None:
        if next_element.tag.endswith("p"):
            return Paragraph(next_element, paragraph._parent)
        if next_element.tag.endswith("tbl"):
            next_element = next_element.getnext()
            continue
        next_element = next_element.getnext()
    return None


def split_pre_section(pre: list[str]) -> tuple[list[str], list[str]]:
    for idx, text in enumerate(pre):
        if text == RELEVANT_DETAILS_HEADING:
            return pre[:idx], pre[idx:]
    legacy_split = 7
    if len(pre) > legacy_split and pre[legacy_split].startswith("This section uses"):
        return pre[:legacy_split], [RELEVANT_DETAILS_HEADING, *pre[legacy_split:]]
    return pre[:legacy_split], pre[legacy_split:]


def sync_data_sources_block(doc: Document, data_pre: list[str]) -> None:
    if not data_pre:
        return

    heading = data_pre[0] if data_pre[0] == DATA_SOURCES_HEADING else None
    body = data_pre[1:] if heading else data_pre

    anchor: Paragraph | None = None
    in_data_sources = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph_is_gray(paragraph) and text == "Data sources":
            in_data_sources = True
            continue
        if not in_data_sources:
            continue
        if paragraph_is_gray(paragraph) and text.startswith("Selection and demographic"):
            break
        if paragraph_is_gray(paragraph) and text == "Note appendices for details.":
            anchor = paragraph

    if anchor is None:
        return

    to_remove: list[Paragraph] = []
    in_data_sources = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph_is_gray(paragraph) and text == "Data sources":
            in_data_sources = True
            continue
        if not in_data_sources:
            continue
        if paragraph_is_gray(paragraph) and text.startswith("Selection and demographic"):
            break
        if not paragraph_is_gray(paragraph) and text:
            to_remove.append(paragraph)

    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)

    current = anchor
    if heading:
        current = insert_paragraph_after(current, heading, style="Heading 3")
    for text in body:
        current = insert_paragraph_after(current, text, style="Normal")


def split_md_sections(black_paragraphs: list[str]) -> tuple[list[str], list[str], list[str]]:
    pre: list[str] = []
    results: list[str] = []
    appendix: list[str] = []
    section = "pre"
    for text in black_paragraphs:
        normalized = text.lstrip("# ").strip()
        if normalized.startswith(RESULTS_START) or normalized == REVENUE_COMPARISON_HEADING:
            section = "results"
        elif normalized.startswith(APPENDIX_START):
            section = "appendix"
            text = normalized
        if section == "pre":
            pre.append(text)
        elif section == "results":
            results.append(normalized if normalized == REVENUE_COMPARISON_HEADING else text)
        else:
            appendix.append(text)
    return pre, results, appendix


def remove_stale_pre_paragraphs(doc: Document) -> None:
    to_remove: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(RESULTS_START) or text.startswith(APPENDIX_START):
            break
        if paragraph_is_gray(paragraph) or paragraph_has_image(paragraph) or not text:
            continue
        if any(text.startswith(prefix) for prefix in STALE_PRE_PREFIXES):
            to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)


def find_appendix_paragraph_index(doc: Document) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(APPENDIX_START):
            return idx
        if text.startswith("The tables below list all p-values"):
            return idx
    raise RuntimeError("Could not find appendix section in document.")


def find_gray_conclusions_paragraph(doc: Document) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph_is_gray(paragraph) and paragraph.text.strip() == "Conclusions":
            return paragraph
    raise RuntimeError("Could not find the gray Conclusions heading.")


def ensure_canonical_appendix_heading(doc: Document) -> Paragraph:
    conclusions = find_gray_conclusions_paragraph(doc)
    conclusions_idx = paragraph_index(doc, conclusions)
    headings = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith(APPENDIX_START)
    ]
    canonical = next(
        (paragraph for paragraph in headings if paragraph_index(doc, paragraph) > conclusions_idx),
        None,
    )

    for paragraph in reversed(headings):
        if canonical is None or paragraph._element is not canonical._element:
            remove_paragraph(paragraph)

    if canonical is not None:
        set_paragraph_text_black(canonical, APPENDIX_START, style="Heading 2")
        return canonical

    paragraphs = doc.paragraphs
    conclusions_idx = paragraph_index(doc, conclusions)
    anchor = conclusions
    for candidate in paragraphs[conclusions_idx + 1 :]:
        if not paragraph_is_gray(candidate):
            break
        anchor = candidate
    return insert_paragraph_after(anchor, APPENDIX_START, style="Heading 2")


def find_nonprofit_results_anchor(doc: Document) -> Paragraph:
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if not paragraph_is_gray(paragraph):
            continue
        if "Results of non-profit organization analysis" not in paragraph.text:
            continue
        for j in range(idx + 1, min(idx + 6, len(paragraphs))):
            candidate = paragraphs[j]
            if paragraph_is_gray(candidate) and "Main analysis content goes here" in candidate.text:
                return candidate

        anchor = paragraph
        for j in range(idx + 1, len(paragraphs)):
            candidate = paragraphs[j]
            text = candidate.text.strip()
            if text.startswith(APPENDIX_START):
                break
            if text == REVENUE_COMPARISON_HEADING or text.startswith(RESULTS_START):
                return anchor
            if paragraph_is_gray(candidate) or paragraph_has_image(candidate):
                anchor = candidate
        raise RuntimeError(
            "Could not find anchor before the non-profit revenue comparison block."
        )
    raise RuntimeError("Could not find non-profit results anchor paragraph.")


def paragraph_index(doc: Document, paragraph: Paragraph) -> int:
    for idx, candidate in enumerate(doc.paragraphs):
        if candidate._element is paragraph._element:
            return idx
    raise RuntimeError("Paragraph not found in document.")


def sync_nonprofit_relevant_details(doc: Document, relevant_pre: list[str]) -> None:
    if not relevant_pre:
        return

    heading = relevant_pre[0] if relevant_pre[0] == RELEVANT_DETAILS_HEADING else None
    details = relevant_pre[1:] if heading else relevant_pre

    target_reader: Paragraph | None = None
    target_results: Paragraph | None = None
    in_nonprofit_section = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph_is_gray(paragraph) and "Characterization of non-profit organizations" in text:
            if "..." not in text:
                in_nonprofit_section = True
            continue
        if not in_nonprofit_section:
            continue
        if paragraph_is_gray(paragraph) and text.startswith("Anything the reader needs to know"):
            target_reader = paragraph
        if paragraph_is_gray(paragraph) and "Results of non-profit organization analysis" in text:
            target_results = paragraph
            break

    if target_reader is None or target_results is None:
        return

    to_remove: list[Paragraph] = []
    removing = False
    for paragraph in doc.paragraphs:
        if paragraph._element is target_reader._element:
            removing = True
            continue
        if paragraph._element is target_results._element:
            break
        if removing and not paragraph_is_gray(paragraph):
            to_remove.append(paragraph)

    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)

    current = target_reader
    if heading:
        current = insert_paragraph_after(current, heading, style="Heading 3")
    for text in details:
        current = insert_paragraph_after(current, text, style=style_for_black_text(text))

def replace_results_block(doc: Document, results: list[str]) -> None:
    conclusions_idx = paragraph_index(doc, find_gray_conclusions_paragraph(doc))
    anchor = find_nonprofit_results_anchor(doc)
    anchor_idx = paragraph_index(doc, anchor)
    to_remove: list[Paragraph] = []
    images_to_remove: list[Paragraph] = []

    for i, paragraph in enumerate(doc.paragraphs):
        if i >= conclusions_idx:
            break
        text = paragraph.text.strip()
        if paragraph_is_gray(paragraph):
            continue
        if i < anchor_idx and (
            text.startswith(RESULTS_START)
            or text == REVENUE_COMPARISON_HEADING
            or text.startswith("Each revenue source was analyzed")
            or text.startswith("Figure 4.")
            or text.startswith("Findings where Black Hills differed")
            or text.startswith("What the 2022 revenue-source comparison shows")
        ):
            to_remove.append(paragraph)
            continue
        if i <= anchor_idx:
            continue
        if paragraph_has_image(paragraph):
            images_to_remove.append(paragraph)
            continue
        if text:
            to_remove.append(paragraph)

    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)
    for paragraph in reversed(images_to_remove):
        remove_image_paragraph(paragraph)

    current = anchor
    for text in results:
        current = insert_paragraph_after(
            current,
            text,
            style=style_for_black_text(text),
        )

def format_p_value(p_value: float) -> str:
    if pd.isna(p_value):
        return ""
    if float(p_value) < 0.001:
        return "< 0.001"
    return f"{float(p_value):.3f}"


def format_dollar(value: float) -> str:
    if float(value) < 0:
        return f"-${abs(float(value)):,.0f}"
    return f"${float(value):,.0f}"


def format_gap_ci(row: pd.Series) -> str:
    gap = row["median_difference"]
    low = row["median_difference_ci_low"]
    high = row["median_difference_ci_high"]
    return (
        f"{format_dollar(gap)} (95% CI {format_dollar(low)} to {format_dollar(high)})"
    )


def format_direction(row: pd.Series) -> str:
    direction = str(row["direction"]).strip()
    suffix = "significant" if float(row["p_value"]) < 0.05 else "not significant"
    if direction.endswith("; significant") or direction.endswith("; not significant"):
        return direction
    return f"{direction}; {suffix}"


def load_significant_results_table_rows() -> list[list[str]]:
    """Significant 2022 pairwise rows in presentation table order."""

    df = pd.read_csv(PAIRWISE_CSV)
    df = df.loc[df["p_value"].lt(0.05)].copy()
    df["variable_order"] = df["variable"].map({value: index for index, value in enumerate(REPORT_VARIABLES)})
    df["region_order"] = df["benchmark_region"].map(
        {value: index for index, value in enumerate(REPORT_REGION_ORDER)}
    )
    df = df.sort_values(["variable_order", "region_order"])
    rows: list[list[str]] = []
    for _, match in df.iterrows():
        rows.append(
            [
                REPORT_DISPLAY_LABELS[str(match["variable"])],
                str(match["benchmark_region"]),
                format_direction(match),
                f"${float(match['black_hills_positive_median']):,.0f}",
                f"${float(match['benchmark_positive_median']):,.0f}",
                format_gap_ci(match),
                format_p_value(float(match["p_value"])),
            ]
        )
    return rows


def find_appendix_template_table(doc: Document) -> Table:
    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header in (
            APPENDIX_TABLE_HEADERS,
            SOURCE_TABLE_HEADERS,
            PRIOR_SOURCE_TABLE_HEADERS,
            LEGACY_SOURCE_TABLE_HEADERS,
        ):
            return table
    raise RuntimeError("Could not find appendix table template in document.")


def copy_cell_format(source_cell, target_cell) -> None:
    src_tc = source_cell._tc
    tgt_tc = target_cell._tc
    src_tcPr = src_tc.find(qn("w:tcPr"))
    if src_tcPr is None:
        return
    existing = tgt_tc.find(qn("w:tcPr"))
    if existing is not None:
        tgt_tc.remove(existing)
    tgt_tc.insert(0, deepcopy(src_tcPr))


def set_cell_text_like_template(target_cell, text: str, template_cell) -> None:
    copy_cell_format(template_cell, target_cell)
    target_cell.text = ""
    paragraph = target_cell.paragraphs[0]
    template_paragraph = template_cell.paragraphs[0]
    paragraph.alignment = template_paragraph.alignment
    run = paragraph.add_run(text)
    if template_paragraph.runs:
        template_run = template_paragraph.runs[0]
        template_r_pr = template_run._element.find(qn("w:rPr"))
        if template_r_pr is not None:
            existing = run._element.find(qn("w:rPr"))
            if existing is not None:
                run._element.remove(existing)
            run._element.insert(0, deepcopy(template_r_pr))


def insert_significant_results_table_after(
    doc: Document,
    paragraph: Paragraph,
    headers: list[str],
    rows: list[list[str]],
) -> Table:
    template = find_appendix_template_table(doc)
    template_header = template.rows[0].cells
    template_data = template.rows[1].cells
    style_map = [1, 0, 1, 2, 3, 4, 5]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = template.style

    all_rows = [headers, *rows]
    for row_idx, row_values in enumerate(all_rows):
        template_row = template_header if row_idx == 0 else template_data
        for col_idx, value in enumerate(row_values):
            set_cell_text_like_template(
                table.rows[row_idx].cells[col_idx],
                value,
                template_row[style_map[col_idx]],
            )

    tbl_element = table._tbl
    doc.element.body.remove(tbl_element)
    paragraph._element.addnext(tbl_element)
    return table


def remove_following_significant_table(paragraph: Paragraph) -> None:
    next_element = paragraph._element.getnext()
    while next_element is not None and next_element.tag.endswith("tbl"):
        parent = paragraph._element.getparent()
        parent.remove(next_element)
        next_element = paragraph._element.getnext()


def upsert_significant_results_table(doc: Document) -> None:
    intro_para: Paragraph | None = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(FINDINGS_INTRO_PREFIX):
            intro_para = paragraph
            break
    if intro_para is None:
        return

    remove_following_significant_table(intro_para)

    to_remove: list[Paragraph] = []
    found_intro = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(FINDINGS_INTRO_PREFIX):
            found_intro = True
            continue
        if not found_intro:
            continue
        if text.startswith(FINDINGS_AFTER_PREFIX):
            break
        if text.startswith("- "):
            to_remove.append(paragraph)
            continue
        if next_element_is_table(paragraph):
            break

    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)

    insert_significant_results_table_after(
        doc,
        intro_para,
        SIGNIFICANT_TABLE_HEADERS,
        load_significant_results_table_rows(),
    )


def next_element_is_table(paragraph: Paragraph) -> bool:
    next_element = paragraph._element.getnext()
    return next_element is not None and next_element.tag.endswith("tbl")


def extract_figure_captions_from_md(md_text: str) -> list[str]:
    captions: list[str] = []
    for line in md_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("**Figure "):
            captions.append(stripped.replace("**", "").strip())
    return captions


def apply_figure_captions(doc: Document, captions: list[str]) -> None:
    """Set the paragraph immediately after each embedded image to the figure caption."""

    caption_idx = 0
    paragraphs = doc.paragraphs
    for i, paragraph in enumerate(paragraphs):
        if not paragraph_has_image(paragraph):
            continue
        if caption_idx >= len(captions):
            break
        next_idx = i + 1
        while next_idx < len(paragraphs) and paragraph_has_image(paragraphs[next_idx]):
            next_idx += 1
        if next_idx >= len(paragraphs):
            break
        caption_para = paragraphs[next_idx]
        if caption_para.text.strip().startswith(SIGNIFICANT_TABLE_HEADERS[0]):
            caption_para = insert_paragraph_after(paragraph, captions[caption_idx])
        elif paragraph_is_gray(caption_para):
            caption_para = insert_paragraph_after(paragraph, captions[caption_idx])
        else:
            set_paragraph_text_black(
                caption_para,
                captions[caption_idx],
                style="Normal",
            )
        caption_idx += 1


def format_report_run(run, *, bold: bool = False, size_pt: int = BODY_FONT_SIZE_PT) -> None:
    run.font.name = "Arial"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = False
    run.underline = False


def clear_paragraph_borders(paragraph: Paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def format_added_paragraph(paragraph: Paragraph) -> None:
    text = paragraph.text.strip()
    is_title = bool(re.match(r"^(Figure|Table) 5-\d+:", text))
    is_heading = paragraph.style.name.startswith("Heading") or text in HEADING_4_TITLES
    for run in paragraph.runs:
        size_pt = HEADER_FONT_SIZE_PT if is_heading else BODY_FONT_SIZE_PT
        format_report_run(
            run,
            bold=is_title or is_heading or bool(run.bold),
            size_pt=size_pt,
        )
    if is_title:
        clear_paragraph_borders(paragraph)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)


def format_added_table(table: Table) -> None:
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            if row_index == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = tc_pr.find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    tc_pr.append(shading)
                shading.set(qn("w:fill"), ADDED_TABLE_HEADER_FILL)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    format_report_run(run, bold=row_index == 0)


def replace_added_image_blobs(doc: Document) -> None:
    image_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph_has_image(paragraph)]
    if len(doc.inline_shapes) != 30:
        raise RuntimeError(f"Expected 30 embedded report images, found {len(doc.inline_shapes)}")
    if len(image_paragraphs) < 10:
        raise RuntimeError(f"Expected at least 10 image paragraphs, found {len(image_paragraphs)}")
    for paragraph, image_path in zip(image_paragraphs[-10:], REPORT_IMAGE_PATHS):
        if not image_path.exists():
            raise FileNotFoundError(image_path)
        blips = paragraph._element.findall(".//" + qn("a:blip"))
        if len(blips) != 1:
            raise RuntimeError(f"Expected one image relationship in paragraph for {image_path.name}")
        relationship_id = blips[0].get(qn("r:embed"))
        image_part = doc.part.related_parts[relationship_id]
        image_part._blob = image_path.read_bytes()


def clear_image_paragraph_text(paragraph: Paragraph) -> None:
    for text_node in paragraph._element.findall(".//" + qn("w:t")):
        text_node.getparent().remove(text_node)


def remove_q9_layout_paragraphs(doc: Document) -> None:
    legacy_headings = {
        "Total revenue",
        "Program service revenue",
        "Total contributions",
        "Government grants received",
        "Federated campaign contributions",
        "Related organization contributions",
        "Membership dues",
        "Fundraising event contributions",
        "Mixed / unclassified contributions",
        "Other revenue",
    }
    report_titles = {*FIGURE_TITLES.values(), *TABLE_TITLES.values()}
    to_remove: list[Paragraph] = []
    in_appendix = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(APPENDIX_START):
            in_appendix = True
            continue
        if paragraph_is_gray(paragraph) or paragraph_has_image(paragraph):
            continue
        if in_appendix:
            to_remove.append(paragraph)
        elif text.startswith("Figure 4.") or text in report_titles:
            to_remove.append(paragraph)
        elif text in legacy_headings:
            to_remove.append(paragraph)
    for paragraph in reversed(to_remove):
        remove_paragraph(paragraph)


def remove_old_q9_value_tables(doc: Document) -> None:
    for table in list(doc.tables):
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if (
            header == APPENDIX_TABLE_HEADERS
            or header == OVERVIEW_TABLE_HEADERS
            or header == SOURCE_TABLE_HEADERS
            or header == PRIOR_SOURCE_TABLE_HEADERS
            or header == LEGACY_SOURCE_TABLE_HEADERS
        ):
            table._element.getparent().remove(table._element)


def insert_report_table_after(
    doc: Document,
    paragraph: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    fractions: tuple[float, ...],
) -> Table:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for row_index, values in enumerate([headers, *rows]):
        for column_index, value in enumerate(values):
            cell = table.rows[row_index].cells[column_index]
            cell.text = str(value)
    table_element = table._tbl
    doc.element.body.remove(table_element)
    paragraph._element.addnext(table_element)
    set_table_full_page_width(table, doc, fractions)
    format_added_table(table)
    return table


def insert_paragraph_after_table(table: Table, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    paragraph = Paragraph(new_p, table._parent)
    set_paragraph_text_black(paragraph, text, style="Normal")
    return paragraph


def report_values() -> tuple[pd.DataFrame, pd.DataFrame]:
    summary = pd.read_csv(PAIRWISE_CSV)
    rows: list[dict[str, object]] = []
    for variable in REPORT_VARIABLES:
        variable_rows = summary.loc[summary["variable"].eq(variable)]
        first = variable_rows.iloc[0]
        rows.append(
            {
                "variable": variable,
                "region": "Black Hills",
                "positive_n": int(first["black_hills_positive_n"]),
                "median": float(first["black_hills_positive_median"]),
                "reporting_rate": float(first["black_hills_nonzero_percent"]),
            }
        )
        for region in REPORT_REGION_ORDER[1:]:
            row = variable_rows.loc[variable_rows["benchmark_region"].eq(region)].iloc[0]
            rows.append(
                {
                    "variable": variable,
                    "region": region,
                    "positive_n": int(row["benchmark_positive_n"]),
                    "median": float(row["benchmark_positive_median"]),
                    "reporting_rate": float(row["benchmark_nonzero_percent"]),
                }
            )
    return summary, pd.DataFrame(rows)


def load_overview_rows(values: pd.DataFrame) -> list[list[str]]:
    rows: list[list[str]] = []
    for variable in REPORT_VARIABLES:
        subset = values.loc[values["variable"].eq(variable)].set_index("region").reindex(REPORT_REGION_ORDER)
        rows.append(
            [
                REPORT_DISPLAY_LABELS[variable],
                *[f"${float(subset.loc[region, 'median']):,.0f}" for region in REPORT_REGION_ORDER],
            ]
        )
    return rows


def load_aggregate_rows() -> list[list[str]]:
    frame = pd.read_csv(RAW_LEVEL_CSV)
    rows: list[list[str]] = []
    for variable in REPORT_VARIABLES:
        subset = frame.loc[frame["variable"].eq(variable)].set_index("region_label").reindex(
            REPORT_REGION_ORDER
        )
        rows.append(
            [
                REPORT_DISPLAY_LABELS[variable],
                *[
                    f"${round(float(subset.loc[region, 'mean']) * float(subset.loc[region, 'n'])):,.0f}"
                    for region in REPORT_REGION_ORDER
                ],
            ]
        )
    return rows


def load_source_rows(summary: pd.DataFrame, values: pd.DataFrame, variable: str) -> list[list[str]]:
    variable_summary = summary.loc[summary["variable"].eq(variable)].set_index("benchmark_region")
    variable_values = values.loc[values["variable"].eq(variable)].set_index("region")
    rows = [
        [
            "Black Hills",
            f"{100 * float(variable_values.loc['Black Hills', 'reporting_rate']):.1f}%",
            f"{int(variable_values.loc['Black Hills', 'positive_n']):,}",
            f"${float(variable_values.loc['Black Hills', 'median']):,.0f}",
            "NA",
            "NA",
        ]
    ]
    for region in REPORT_REGION_ORDER[1:]:
        row = variable_summary.loc[region]
        rows.append(
            [
                region,
                f"{100 * float(variable_values.loc[region, 'reporting_rate']):.1f}%",
                f"{int(variable_values.loc[region, 'positive_n']):,}",
                f"${float(variable_values.loc[region, 'median']):,.0f}",
                format_gap_ci(row),
                format_p_value(float(row["p_value"])),
            ]
        )
    return rows


def gray_teammate_footnotes(doc: Document) -> None:
    target_prefixes = (
        "State tax provisions reflect rules in effect for the 2022 tax year",
        "Both r and rs are correlation coefficients",
    )
    footnotes_part = next(
        (
            part
            for part in doc.part.package.parts
            if str(part.partname) == "/word/footnotes.xml"
        ),
        None,
    )
    if footnotes_part is None:
        raise RuntimeError("Could not locate footnotes.xml")

    root = etree.fromstring(footnotes_part.blob)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    matched: set[str] = set()
    for footnote in root.xpath("//w:footnote", namespaces=namespace):
        text = "".join(footnote.xpath(".//w:t/text()", namespaces=namespace)).strip()
        prefix = next((item for item in target_prefixes if text.startswith(item)), None)
        if prefix is None:
            continue
        matched.add(prefix)
        for run in footnote.xpath(".//w:r", namespaces=namespace):
            run_properties = run.find(qn("w:rPr"))
            if run_properties is None:
                run_properties = etree.Element(qn("w:rPr"))
                run.insert(0, run_properties)
            color = run_properties.find(qn("w:color"))
            if color is None:
                color = etree.Element(qn("w:color"))
                run_properties.append(color)
            color.set(qn("w:val"), GRAY)

    missing = set(target_prefixes) - matched
    if missing:
        raise RuntimeError(f"Could not locate teammate footnotes: {sorted(missing)}")
    footnotes_part._blob = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )


def rebuild_q9_figures_and_tables(doc: Document) -> None:
    replace_added_image_blobs(doc)
    remove_q9_layout_paragraphs(doc)
    remove_old_q9_value_tables(doc)
    summary, values = report_values()
    image_paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph_has_image(paragraph)]
    if len(image_paragraphs) < 10:
        raise RuntimeError(f"Expected at least 10 Q9 image paragraphs, found {len(image_paragraphs)}")
    appendix_images = image_paragraphs[-10:]

    appendix_heading = next(
        paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith(APPENDIX_START)
    )
    appendix_intro = insert_paragraph_after(appendix_heading, APPENDIX_INTRO, style="Normal")
    format_added_paragraph(appendix_intro)
    overview_table_title = insert_paragraph_after(
        appendix_intro,
        TABLE_TITLES["overview"],
        style="Normal",
    )
    overview_table = insert_report_table_after(
        doc,
        overview_table_title,
        OVERVIEW_TABLE_HEADERS,
        load_overview_rows(values),
        (0.28, 0.144, 0.144, 0.144, 0.144, 0.144),
    )
    aggregate_note = insert_paragraph_after_table(overview_table, AGGREGATE_TABLE_NOTE)
    format_added_paragraph(aggregate_note)
    aggregate_table_title = insert_paragraph_after(
        aggregate_note,
        TABLE_TITLES["aggregate"],
        style="Normal",
    )
    format_added_paragraph(aggregate_table_title)
    insert_report_table_after(
        doc,
        aggregate_table_title,
        OVERVIEW_TABLE_HEADERS,
        load_aggregate_rows(),
        (0.28, 0.144, 0.144, 0.144, 0.144, 0.144),
    )

    last_table: Table | None = None
    for image_paragraph, variable in zip(appendix_images, REPORT_VARIABLES):
        insert_paragraph_before(image_paragraph, FIGURE_TITLES[variable], style="Normal")
        table_title = insert_paragraph_after(image_paragraph, TABLE_TITLES[variable], style="Normal")
        last_table = insert_report_table_after(
            doc,
            table_title,
            SOURCE_TABLE_HEADERS,
            load_source_rows(summary, values, variable),
            (0.12, 0.12, 0.13, 0.16, 0.34, 0.13),
        )
    if last_table is not None:
        definitions_title = insert_paragraph_after_table(last_table, TABLE_TITLES["definitions"])
        definitions_table = insert_report_table_after(
            doc,
            definitions_title,
            REVENUE_DEFINITION_TABLE_HEADERS,
            REVENUE_DEFINITION_TABLE_ROWS,
            REVENUE_DEFINITION_COLUMN_FRACTIONS,
        )
        final_note = insert_paragraph_after_table(definitions_table, APPENDIX_FINAL_NOTE)
        format_added_paragraph(final_note)


def apply_black_paragraph_formatting(doc: Document) -> None:
    """Ensure synced black text uses Normal/Heading styles, not inherited Heading 2."""

    in_appendix = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(APPENDIX_START):
            in_appendix = True
        if paragraph_is_gray(paragraph) or paragraph_has_image(paragraph):
            continue
        if not text:
            continue
        set_paragraph_style(paragraph, style_for_black_text(text, in_appendix=in_appendix))
        format_added_paragraph(paragraph)

    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if (
            header == REVENUE_DEFINITION_TABLE_HEADERS
            or header == SIGNIFICANT_TABLE_HEADERS
            or header == OVERVIEW_TABLE_HEADERS
            or header == SOURCE_TABLE_HEADERS
            or header == LEGACY_SOURCE_TABLE_HEADERS
        ):
            format_added_table(table)


def remove_empty_black_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph_is_gray(paragraph) or paragraph_has_image(paragraph):
            continue
        if not paragraph.text.strip():
            remove_paragraph(paragraph)


def sync_docx(docx_path: Path, black_paragraphs: list[str]) -> Path:
    doc = Document(docx_path)
    remove_aggregate_share_table(doc)
    remove_revenue_definition_tables(doc)
    remove_significant_results_tables(doc)
    pre, results, appendix = split_md_sections(black_paragraphs)

    remove_stale_pre_paragraphs(doc)
    data_pre, relevant_pre = split_pre_section(pre)
    sync_data_sources_block(doc, data_pre)
    sync_nonprofit_relevant_details(doc, relevant_pre)

    replace_results_block(doc, results)

    ensure_canonical_appendix_heading(doc)
    remove_empty_black_paragraphs(doc)
    rebuild_q9_figures_and_tables(doc)
    apply_black_paragraph_formatting(doc)
    gray_teammate_footnotes(doc)

    try:
        doc.save(docx_path)
        return docx_path
    except PermissionError:
        fallback = docx_path.with_name(docx_path.stem + "_synced.docx")
        doc.save(fallback)
        print(f"  (file locked; saved to {fallback.name})")
        return fallback


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    black_paragraphs = extract_black_paragraphs_from_md(md_text)
    pre, results, appendix = split_md_sections(black_paragraphs)
    print(f"Markdown black paragraphs: {len(black_paragraphs)} (pre={len(pre)}, results={len(results)}, appendix={len(appendix)})")
    gray_check = "Need to write it at the end based on conclusions"
    if not DOCX_PATH.exists():
        print(f"Missing: {DOCX_PATH}")
        return
    saved = sync_docx(DOCX_PATH, black_paragraphs)
    doc = Document(saved)
    full = "\n".join(p.text for p in doc.paragraphs)
    sig_table_ok = any(
        [c.text.strip() for c in t.rows[0].cells] == SIGNIFICANT_TABLE_HEADERS
        for t in doc.tables
    )
    figure_count = sum(full.count(f"Figure 5-{index}:") for index in range(2, 12))
    table_count = sum(full.count(f"Table 5-{index}:") for index in range(1, 14))
    print(
        f"{saved.name}: gray intact={gray_check in full}, "
        f"sample-size heading={'How many organizations could we compare?' in full}, "
        f"eligible pool 256={'256 Black Hills organization-years' in full}, "
        f"nonprofit data ref={'described under Data sources' in full}, "
        f"main-body p-value language={'p-value' in full.split('Conclusions', 1)[0]}, "
        f"sig results table removed={not sig_table_ok}, "
        f"aggregate gone={'Black Hills share of aggregate' not in full}, "
        f"figure 5-1 removed={'Figure 5-1:' not in full}, "
        f"figures 5-2 to 5-11={figure_count == 10}, "
        f"tables 5-1 to 5-13={table_count == 13}, "
        f"embedded images={len(doc.inline_shapes)}, "
        f"results body Normal={all(doc.paragraphs[i].style.name == 'Normal' for i in range(len(doc.paragraphs)) if doc.paragraphs[i].text.strip().startswith('Only organizations that reported'))}, "
        f"h3 sample size={any(p.text.strip()=='How many organizations could we compare?' and p.style.name=='Heading 3' for p in doc.paragraphs)}, "
        f"revenue defs table={any([c.text.strip() for c in t.rows[0].cells]==REVENUE_DEFINITION_TABLE_HEADERS for t in doc.tables)}, "
        f"h3 data sources={any(p.text.strip()==DATA_SOURCES_HEADING and p.style.name=='Heading 3' for p in doc.paragraphs)}, "
        f"h3 relevant details={any(p.text.strip()==RELEVANT_DETAILS_HEADING and p.style.name=='Heading 3' for p in doc.paragraphs)}, "
        f"h3 revenue comparison={sum(1 for p in doc.paragraphs if p.text.strip()==REVENUE_COMPARISON_HEADING)==1}"
    )


if __name__ == "__main__":
    main()
